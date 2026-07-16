"""
Factory Audit document matcher + FA template filler.

Input: folder with client QMS documents (and usually a master list PDF).
Output: filled FA template .docx + summary JSON/txt.

Document numbers are looked up per folder — never hardcoded.
"""

from __future__ import annotations

import json
import re
import shutil
import zipfile
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

# Register all standard Word/DrawingML namespace prefixes so ET.tostring
# never emits generic ns0/ns1/... prefixes, which make Word refuse to open files.
for _p, _u in {
    "":         "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w":        "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r":        "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp":       "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "wp14":     "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing",
    "a":        "http://schemas.openxmlformats.org/drawingml/2006/main",
    "a14":      "http://schemas.microsoft.com/office/drawing/2010/main",
    "adp":      "http://schemas.microsoft.com/office/drawing/2014/main",
    "aink":     "http://schemas.microsoft.com/office/drawing/2016/ink",
    "am3d":     "http://schemas.microsoft.com/office/drawing/2017/model3d",
    "pic":      "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "pic14":    "http://schemas.microsoft.com/office/drawing/2010/picture",
    "m":        "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "mc":       "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "oel":      "http://schemas.microsoft.com/office/2019/extlst",
    "wpc":      "http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas",
    "wpg":      "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
    "wpi":      "http://schemas.microsoft.com/office/word/2010/wordprocessingInk",
    "wps":      "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "wne":      "http://schemas.microsoft.com/office/word/2006/wordml",
    "w14":      "http://schemas.microsoft.com/office/word/2010/wordml",
    "w15":      "http://schemas.microsoft.com/office/word/2012/wordml",
    "w16":      "http://schemas.microsoft.com/office/word/2018/wordml",
    "w16cex":   "http://schemas.microsoft.com/office/word/2018/wordml/cex",
    "w16cid":   "http://schemas.microsoft.com/office/word/2016/wordml/cid",
    "w16du":    "http://schemas.microsoft.com/office/word/2023/wordml/word16du",
    "w16sdtdh": "http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash",
    "w16sdtfl": "http://schemas.microsoft.com/office/word/2024/wordml/sdtformatlock",
    "w16se":    "http://schemas.microsoft.com/office/word/2015/wordml/symex",
    "cx":       "http://schemas.microsoft.com/office/drawing/2014/chartex",
    "cx1":      "http://schemas.microsoft.com/office/drawing/2015/9/8/chartex",
    "cx2":      "http://schemas.microsoft.com/office/drawing/2015/10/21/chartex",
    "cx3":      "http://schemas.microsoft.com/office/drawing/2016/5/9/chartex",
    "cx4":      "http://schemas.microsoft.com/office/drawing/2016/5/10/chartex",
    "cx5":      "http://schemas.microsoft.com/office/drawing/2016/5/11/chartex",
    "cx6":      "http://schemas.microsoft.com/office/drawing/2016/5/12/chartex",
    "cx7":      "http://schemas.microsoft.com/office/drawing/2016/5/13/chartex",
    "cx8":      "http://schemas.microsoft.com/office/drawing/2016/5/14/chartex",
    "v":        "urn:schemas-microsoft-com:vml",
    "o":        "urn:schemas-microsoft-com:office:office",
    "xdr":      "http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing",
}.items():
    ET.register_namespace(_p, _u)

from PIL import Image, ImageOps
from pypdf import PdfReader

try:
    import pytesseract
except ImportError:
    pytesseract = None  # type: ignore[assignment]

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
A = f"{{{A_NS}}}"
PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"
PIC = f"{{{PIC_NS}}}"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
WP = f"{{{WP_NS}}}"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
R = f"{{{R_NS}}}"

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp"}
PDF_PHOTO_HINTS = [
    "workshop", "photo", "image", "building", "floor", "area", "plant",
    "yard", "packing", "inspection", "lab", "storage", "warehouse", "front",
]
IMAGE_EXCLUDE_HINTS = [
    "procedure", "ccl-", "-pr-", "-fr-", "inspection control", "trade license",
    "organizational", "iso", "certificate", "application form", "master list",
    "matster list", "itp", "register", "complaint", "ncr", "dwg", "drawing",
    "organizational chart", "factory info", "audit plan", "custom clearance",
    "signature", "application form", "organizational chart",
    "9k", "14k", "45k", "e-cert", "multisite", "_cert_previews",
    "quality plan", "flow chart", "work instruction", "record of the testing",
    "record of the visual", "factory layout",
]
IMAGE_MAX_WIDTH_PX = 640
IMAGE_MAX_WIDTH_EMU = 2_560_000   # ~2.8 inches — fits a 3.25-inch 2-col table cell
PHOTO_FIXED_WIDTH_EMU = 1_800_000  # ~1.97 inches — uniform small size for photo table
PHOTO_FIXED_HEIGHT_EMU = 1_350_000  # ~1.48 inches — keeps photo rows compact and aligned
SIGNATURE_MAX_WIDTH_EMU = 1_350_000  # compact signatures for sign-off cells
SUPPORT_DOC_MAX_WIDTH_EMU = 1_700_000  # ISO/GMP cert previews ~1.85 inches per page
SUPPORT_DOC_MAX_HEIGHT_EMU = 2_200_000  # cap cert page height so images stay under headings
PHOTO_TABLE_FULL_WIDTH = 9360   # dxa ≈ 6.5 inches (text body width)
PHOTO_CELL_WIDTH = 4680         # dxa ≈ 3.25 inches per column
FACTORY_PHOTOS_START = "Factory Photos:"
FACTORY_PHOTOS_END = "Factory Layout"
CERT_HEADING_LABELS = frozenset({"ISO 9001", "ISO 14001", "ISO 45001", "GMP"})
CHECKED = "☒"
UNCHECKED = "☐"
AUDITOR_REGISTRY: dict[str, str] = {
    "PB": "Pranav Bhatnagar",
    "PD": "Pradeep Divakaran",
    "MH": "Mohammed Hashir",
    "NP": "Nigal Paul",
}
MAIN_AUDITOR_CODES = {"PB", "PD", "NP"}
AUDITOR_SIGNATURE_DIR = Path(__file__).resolve().parent / "signatures"

DEFAULT_TEMPLATE = Path.home() / "Desktop" / "FA template.docx"
SCRIPT_DIR = Path(__file__).resolve().parent

MASTER_LIST_PATTERNS = [
    r"master\s*(?:documents?\s*)?list",
    r"list\s*of\s*.*(?:procedures?|documents?|records?)",
    r"list\s*of\s*ims",
    r"matster\s*list",
    r"daftar\s*dokumen",   # Indonesian "list of documents"
]

# Broad doc-number pattern: accepts IMS-English, Indonesian WG/..., S*-WG-*, and alphanumeric codes.
# Any format starting with letters and containing digits is accepted.
DOC_NO_RE = re.compile(
    r"\b("
    r"WG/[A-Z]{2,6}-[A-Z]{2,6}/\d{3,4}"   # WG/PRO-PROD/076
    r"|S[A-Z]{2}-WG-\d{2,4}"               # SMQ-WG-02, SML-WG-01
    r"|[A-Z]{2,8}(?:[/-][A-Z]{2,8}){1,5}-\d{3,4}"  # GPGM-IMS-SOP-001, WG-IK-PROD-001
    r"|[A-Z]{2,8}(?:-[A-Z]{0,6}\d{3,6})+"  # SOP-QC-001, FR-014, QP-001
    r"|[A-Z]{2,8}-\d{4,8}"                 # MS-0042909
    r")\b",
    re.I,
)

# Broader doc-number pattern for filename prefixes: e.g. "1.1", "MS-0042909", "SOP-14", "QP-001"
_FILENAME_DOC_NO_RE = re.compile(
    r"^(\d+\.\d+(?:\.\d+)*"               # 1.1 / 1.1.2
    r"|[A-Z]{1,6}-\d{4,}"                 # MS-0042909, FR-0012
    r"|[A-Z]{2,6}-[A-Z]{2,6}-\w{2,6}-\d{2,4}"  # ABCHIL-QMS-SOP-14
    r"|[A-Z]{2,6}-[A-Z]{2,10}-\d{2,4}"   # QC-INSP-01
    r")\s*",
    re.I,
)

SECTION_IDS = [
    "2.1", "2.2", "2.3",
    "3.1", "3.2", "3.3", "3.4", "3.5", "3.6",
    "4.1", "4.2", "4.3", "4.4", "4.5",
    "5.1", "5.2",
    "6.1", "6.2", "6.3",
]

# Keywords per section — generic, not factory-specific.
SECTION_KEYWORDS: dict[str, list[str]] = {
    # 2.0 Verification of purchased components/materials (incoming)
    "2.1": [  # Procedure for incoming inspection / verify materials upon receipt
        "incoming inspection", "incoming material inspection", "material inspection",
        "material receipt", "receiving", "verify material", "verification of material",
        "report incoming material", "incoming material", "incomming", "sop for incoming",
    ],
    "2.2": [  # Conformity of supplier material / handling non-conforming components
        "conformity", "certificate of conformity", "coa", "mtc", "certificate of analysis",
        "mill test", "supplier certificate", "supplier test", "material certificate",
        "non-conforming", "non conforming", "nonconforming",
    ],
    "2.3": [  # Checklist / records of incoming material inspection
        "incoming inspection record", "incoming record", "incoming checklist",
        "material receiving inspection", "receiving report", "records of the incoming",
        "report incoming material", "inspection checklist", "mir",
        "incoming inspection report", "incoming inspection",
    ],
    # 3.0 Production control, monitoring and testing
    "3.1": [  # Training procedures / records
        "training", "training procedure", "training record", "training plan",
        "competency", "competence", "awareness", "toolbox",
    ],
    "3.2": [  # Preventive maintenance of equipment / work instructions / quality plan of process
        "preventive maintenance", "maintenance plan", "maintenance procedure",
        "equipment management", "work instruction", "quality plan", "process quality",
        "flow chart process", "flow chart", "process control",
    ],
    "3.3": [  # Inspection / testing of in-process products
        "in-process", "in process", "in line inspection", "in-line inspection",
        "inline inspection", "qc in line", "process inspection",
        "inspection and test plan", "itp", "wi qc in line",
        "end table quality", "end-table quality", "end table",
    ],
    "3.4": [  # Inspection/testing SOPs, work instructions, half-way product instructions
        "work instruction", "operator instruction", "process instruction",
        "standard operating", "method statement", "knitting process", "trimming",
        "pairing", "reversing", "heat set", "metal detector", "rosso",
        "instruksi kerja", "ik :",
        "quality control guideline", "quality guideline", "internal quality plan",
        "quality procedure", "qc guideline",
    ],
    "3.5": [  # Non-conforming product control (in process)
        "non-conforming product", "non conforming product", "nonconforming",
        "ncr", "control of non", "non conformity", "rejection tag", "rejection sticker",
    ],
    "3.6": [  # Records of SOP being followed / routine test & maintenance records
        "routine test", "inspection record", "test report", "maintenance record",
        "maintenance log", "record of sop", "production record", "pm plan",
    ],
    # 4.0 Calibration of safety test and measuring equipment
    "4.1": [  # Calibration certificate (number reviewed)
        "calibration certificate", "calibration cert", "certificate calibration",
        "certification", "calibration",
    ],
    "4.2": [  # List of equipment with calibration due dates / logs
        "equipment list", "calibration register", "calibration due", "calibration log",
        "calibration schedule", "equipment master", "measuring equipment", "instrument list",
        "lab profile", "laboratory profile", "dyeing laboratory", "testing laboratory",
    ],
    "4.3": [  # SOP / procedure of calibration / lab accreditation
        "calibration procedure", "calibration sop", "calibration control",
        "calibration method", "procedure of calibration", "sop calibration",
        "lab accreditation", "laboratory accreditation", "in-house lab certificate",
        "in house lab certificate", "in house lab", "lab approval", "accreditation certificate",
        "accrediation certificate",
    ],
    "4.4": [  # Records / testing done for calibration
        "calibration record", "calibration result", "calibration test",
        "validation", "verification record",
    ],
    "4.5": [  # Corrective action for non-conformity / product recall
        "corrective action", "recall", "non conformity", "functional check",
        "unsatisfactory", "capa",
    ],
    # 5.0 Handling and storage
    "5.1": [  # Raw material handling
        "material handling", "raw material handling", "component handling",
        "yarn receipt", "material storage", "storage of material",
        "handling and storage", "warehouse",
    ],
    "5.2": [  # Final product handling
        "finished product", "finished goods", "final product handling",
        "final product storage", "finished goods storage", "packing", "packaging",
        "shipping",
    ],
    # 6.0 Product Verification Testing (PVT) and complaint handling
    "6.1": [  # Final product testing
        "final inspection", "final test", "finished good", "pvt", "product verification",
        "final product test", "record of the testing", "visual inspection",
        "finished product inspection",
    ],
    "6.2": [  # Final product rejection control / product traceability
        "traceability", "product traceability", "rejection control", "rejection",
        "unsatisfactory", "pvt result",
    ],
    "6.3": [  # Customer complaint handling
        "complaint", "customer complaint", "customer related", "complaint handling",
        "complaint log",
    ],
}

FILENAME_HINTS: dict[str, list[str]] = {
    "2.1": ["incoming", "receipt", "receiving", "material inspection", "report incoming"],
    "2.2": ["conformity", "coa", "mtc", "non-conform", "nonconform", "supplier"],
    "2.3": ["incoming record", "incoming checklist", "receiving report", "report incoming", "mir", "incoming inspection"],
    "3.1": ["training", "toolbox", "competenc"],
    "3.2": ["preventive", "maintenance", "work instruction", "wi ", "quality plan", "flow chart"],
    "3.3": ["in line", "in-line", "inline", "qc in line", "itp", "in process", "end table"],
    "3.4": ["sop", "work instruction", "wi ", "knitting", "trimming", "pairing", "yarn", "metal detector", "guideline", "internal quality"],
    "3.5": ["ncr", "non-conform", "nonconform", "rejection"],
    "3.6": ["routine", "maintenance", "test report", "record", "pm plan"],
    "4.1": ["calibration certificate", "calibration cert", "certification", "certificate calibration", "calibration"],
    "4.2": ["equipment list", "calibration register", "calibration log", "calibration due", "calibration schedule", "lab profile", "laboratory"],
    "4.3": ["calibration control", "calibration procedure", "calibration sop", "calibration method", "lab accreditation", "accreditation", "in house lab", "accrediation"],
    "4.4": ["calibration record", "validation", "calibration result"],
    "4.5": ["corrective", "recall", "ncr"],
    "5.1": ["handling", "warehouse", "storage", "yarn receipt"],
    "5.2": ["packing", "shipping", "finished", "final product"],
    "6.1": ["final inspection", "final test", "pvt", "record of the testing", "visual inspection"],
    "6.2": ["traceability", "rejection", "ncr"],
    "6.3": ["complaint"],
}


@dataclass
class PhotoSlot:
    slot_id: str
    label: str
    anchor: str
    anchor_type: str  # paragraph | table_row | section_row
    keywords: list[str]


PHOTO_SLOTS: list[PhotoSlot] = [
    PhotoSlot("1.5_plant", "Manufacturing Plant Outlook", "Manufacturing Plant Outlook", "table_row",
              ["building", "front", "main", "outlook", "plant", "factory"]),
    PhotoSlot("1.5_workshop", "Each Floor/Workshop/Process", "Each Floor/Workshop/Process", "table_row",
              ["workshop", "floor", "production", "fabrication", "yard", "shop"]),
    PhotoSlot("1.5_warehouse", "Material and Final Product Warehouse", "Material and Final Product Warehouse", "table_row",
              ["warehouse", "storage", "material"]),
    PhotoSlot("1.5_audit_loc", "Audit Equipment and Location", "Audit Equipment and Location", "table_row",
              ["audit", "lab", "inspection area", "equipment location"]),
    PhotoSlot("1.5_product", "Product / packaging / trademark photo", "Photographs of the actual product", "table_row",
              ["product", "packaging", "trademark", "brand", "label"]),
    PhotoSlot("fp_front", "Front Main Building", "Front Main Building", "paragraph",
              ["front", "main", "building", "outlook", "gate"]),
    PhotoSlot("fp_incoming", "Incoming Raw Material Area", "Incoming Raw Material", "paragraph",
              ["incoming", "raw material", "rm area", "receiving"]),
    PhotoSlot("fp_production", "Production floor Area", "Production floor", "paragraph",
              ["production", "workshop", "fabrication", "floor", "yard", "shop"]),
    PhotoSlot("fp_packing", "Packing Area", "Packing Area", "paragraph",
              ["packing", "packaging"]),
    PhotoSlot("fp_inspection", "Inspection Area", "Inspection Area", "paragraph",
              ["inspection area", "qc area", "quality"]),
    PhotoSlot("fp_lab", "Lab Area", "Lab Area Equipment", "paragraph",
              ["lab area", "laboratory", "testing equipment", "equipment used for testing", "equipment used"]),
    PhotoSlot("fp_cal_label", "Calibration Labels on Equipment", "Calibration Labels on Equipment", "paragraph",
              ["calibration label", "cal label", "calibration tag", "calibration label on"]),
    PhotoSlot("fp_rejection", "Non-Conformity / Rejection Area", "Non-Conformity Area", "paragraph",
              ["rejection area", "rejection sticker", "non-conformity area", "nonconform", "ncr",
               "non comformity", "non conformity"]),
    PhotoSlot("fp_final_storage", "Final Product Storage Area", "Final Product Storage Area", "paragraph",
              ["final product storage", "finished storage", "fg storage", "storage area"]),
    PhotoSlot("fp_product_label", "Product label image", "product label", "paragraph",
              ["product label", "image of the product label"]),
    PhotoSlot("fp_barcode", "Barcode / QR code", "barcode", "paragraph",
              ["barcode", "qr code", "qr", "product barcode"]),
    PhotoSlot("fp_product_shape", "Product identity / shape", "identity and the shape", "paragraph",
              ["product photo", "identity", "shape", "model", "final product photo"]),
    PhotoSlot("fp_trademark", "Trademark image", "trademark", "paragraph",
              ["trademark", "trade mark", "brand logo", "logo"]),
    PhotoSlot("sec_4.2", "Calibration label (section 4.2)", "4.2", "section_row",
              ["calibration label", "cal label"]),
]

# Deterministic ordered layout of the Factory Photos section.
# (heading text as it appears in the report, source key for its images)
FACTORY_PHOTO_LAYOUT: list[tuple[str, str]] = [
    ("Front Main Building", "fp_front"),
    ("Incoming Raw Material (RM) Area", "fp_incoming"),
    ("Production floor Area", "fp_production"),
    ("Packing Area", "fp_packing"),
    ("Inspection Area", "fp_inspection"),
    ("Lab Area Equipment used for Testing and inspection", "fp_lab"),
    ("Calibration Labels on Equipment", "fp_cal_label"),
    ("Non-Conformity Area/Rejection Area for RM & Final Product", "fp_rejection"),
    ("Final Product Storage Area", "fp_final_storage"),
    ("An image of the product label", "fp_product_label"),
    ("An image of the product barcode / QR code", "fp_barcode"),
    ("An image shows the identity and the shape of the product, images of each model", "fp_product_shape"),
    ("An image of the trademark", "fp_trademark"),
    ("ISO 9001", "iso9001"),
    ("ISO 14001", "iso14001"),
    ("Factory Layout", "factory_layout"),
    ("Master list for the Documents - Procedures & Records", "master_list"),
    ("Organization Structure", "org_structure"),
]


def detect_general_doc_images(folder: Path) -> dict[str, Path]:
    """Locate the single best file for Factory Layout, Master list and Org chart.

    These appear as image slots in the Factory Photos section and must show only
    their own document (not random factory photos).
    """
    specs = {
        "factory_layout": [r"factory\s*layout", r"\blayout\b", r"floor\s*plan", r"site\s*plan"],
        "master_list": [r"master\s*list", r"list of documents", r"document.*register"],
        "org_structure": [r"organi[sz]ation\s*chart", r"org\s*chart",
                          r"organi[sz]ation\s*structure", r"organi[sz]ational\s*chart"],
    }
    found: dict[str, Path] = {}
    files = sorted(
        (p for p in folder.rglob("*") if p.is_file()),
        # Prefer shallower paths (top-level general documents over nested copies)
        key=lambda p: (len(p.parts), p.name.lower()),
    )
    for key, patterns in specs.items():
        for p in files:
            name = p.name.lower()
            if p.suffix.lower() not in (IMAGE_EXTENSIONS | {".pdf"}):
                continue
            if any(re.search(pat, name) for pat in patterns):
                found[key] = p
                break
    return found

# Full FA checklist — (section title, [(item label, match keywords, optional photo_slot_id)])
CHECKLIST_SECTIONS: list[tuple[str, list[tuple[str, list[str], str | None]]]] = [
    ("GENERAL DOCUMENTS", [
        ("Trade License / Business license / Manufacturing license",
         ["trade license", "business license", "manufacturing license",
          "incorporation", "companies house", "articles of association",
          "commercial registration", "cr certificate", "trade registration",
          "nib", "business license_nib"], None),
        ("ISO 9001", ["iso 9001", "iso9001", "9001 multisite", "9k multisite", "iso9001-2015", "iso 9001-2015"], None),
        ("ISO 14001 (if available)", ["iso 14001", "iso14001", "14k multisite", "iso 14001-2015"], None),
        ("Factory Layout", ["factory layout", "layout", "general arrangement", "dwg-ff", "dwg ff",
                             "floor plan", "site plan", "plant layout", "warehouse layout", "wh1"], None),
        ("Master list for the Documents - Procedures & Records",
         ["master list", "matster list", "list of procedures", "list of ims",
          "quality manual", "ims en", "document control register", "dcr", "ims manual",
          "master list of documents"], None),
        ("Organization Structure",
         ["organizational chart", "organization chart", "org chart", "organization structure",
          "organisation chart", "organisation structure"], None),
    ]),
    ("INCOMING MATERIAL CONTROL", [
        ("Quality Plan – Incoming", ["incoming quality plan", "quality plan incoming",
                                     "raw material control", "material control plan",
                                     "quality plan", "incomming"], None),
        ("COA / MTC / Test Report from the Supplier",
         ["coa", "mtc", "mill test", "certificate of analysis", "supplier test",
          "cof a", "certificate of conformity", "material certificate", "sf-016", "sf016",
          "ba e4", "batch certificate", "batch report"], None),
        ("SOP for Incoming Material Inspection",
         ["incoming material inspection", "material receipt inspection", "material inspection control",
          "incoming inspection", "traceability and inspection", "formulation",
          "report incoming material", "incoming material"], None),
        ("Incoming Inspection Checklist", ["incoming inspection checklist", "incoming checklist",
                                           "receiving checklist", "report incoming material"], None),
        ("Recent Incoming Inspection Records",
         ["incoming inspection record", "material receiving inspection", "mir", "receiving report",
          "raw material control audit", "formulation audit", "incoming record",
          "report incoming material"], None),
        ("Non-Conformance Handling Procedure",
         ["non-conforming", "non conforming", "nonconforming", "non conformity handling",
          "control of non", "ncr procedure", "non-conformance"], None),
        ("Material Rejection Tag", ["rejection tag", "rejection sticker", "material rejection"], None),
    ]),
    ("PROCESS CONTROL", [
        ("Quality Plan (In-process) – Inspection / Testing",
         ["inspection and test plan", "itp", "in-process", "in process quality", "quality plan",
          "quality control plan", "control plan", "qc plan", "process quality"], None),
        ("Training Plan / Procedure",
         ["training plan", "training procedure", "competency", "training and awareness"], None),
        ("Training Records", ["training record", "toolbox talk", "toolbox", "training log"], None),
        ("Work Instructions / Engineering Drawings",
         ["work instruction", "engineering drawing", "dwg", "drawing", "design",
          "operator instruction", "operator work instruction", "process map",
          "work order", "method statement", " wi ", "wi knitting", "wi qc", "wi packing"], None),
        ("Preventive Maintenance SOP",
         ["preventive maintenance procedure", "maintenance procedure", "equipment management",
          "maintenance sop"], None),
        ("Equipment Maintenance Schedule / Log / Plan",
         ["maintenance schedule", "maintenance plan", "maintenance log", "preventive maintenance plan",
          "equipment maintenance"], None),
    ]),
    ("CALIBRATION OF MEASURING EQUIPMENT", [
        ("Calibration SOP", ["calibration control", "calibration procedure", "calibration sop",
                              "calibration check", "calibration method"], None),
        ("Calibration Certificates", ["calibration certificate", "calibration cert",
                                       "calibration check", "instrument calibration",
                                       "calibration of instrument", "certificate calibration",
                                       "certification crock", "certification moisture",
                                       "certification metal", "certification timbangan"], None),
        ("Equipment Master List", ["calibration register", "equipment master", "equipment list",
                                    "measuring equipment", "instrument list", "instruments"], None),
        ("Calibration Tags / Labels", ["calibration label", "calibration tag", "cal label",
                                        "calibration label on the"], None),
    ]),
    ("100% INSPECTION OF FINISHED PRODUCT", [
        ("Quality Plan – Finished Good (testing / inspection)",
         ["finished good", "final inspection", "finished product inspection", "itp",
          "quality control sampling", "qc sampling", "sampling record",
          "flow chart finished", "quality plan"], None),
        ("Records of the Testing", ["test report", "test record", "testing record",
                                     "pneumatic test", "hydrostatic", "qc sheet",
                                     "testing log", "inspection log", "inspection testing",
                                     "quality control log", "record of the testing"], None),
        ("Records of the Visual Inspection", ["visual inspection", "inspection report",
                                               "general inspection", "qc inspection",
                                               "quality control record", "sampling record",
                                               "quality record", "record of the visual"], None),
    ]),
    ("MATERIAL HANDLING & STORAGE", [
        ("Material Handling & Storage SOP", ["material handling", "warehouse management",
                                              "storage procedure", "handling procedure"], None),
        ("Finished Goods Storage Procedure", ["finished goods storage", "finished product storage",
                                               "finished goods", "fg storage"], None),
        ("Packaging / Handling Instructions / Procedures",
         ["packing procedure", "packaging", "shipping procedure", "handling instruction",
          "packing instruction"], None),
    ]),
    ("TEST REPORTS & COMPLAINTS", [
        ("Test Reports", ["test report", "mir", "inspection report", "qc sheet",
                           "testing log", "inspection log"], None),
        ("Customer Complaint Handling SOP",
         ["customer complaint", "complaint handling", "customer related", "complaint register"], None),
        ("Complaint Log", ["complaint register", "complaint log"], None),
        ("Complaint and Corrective ACTION Report",
         ["corrective action report", "complaint and corrective", "car report",
          "corrective action", "change control"], None),
    ]),
    ("FACTORY PHOTOS", [
        ("Front Main Building", ["front", "main building", "building"], "fp_front"),
        ("Incoming Raw Material (RM) Area", ["incoming", "raw material", "rm area"], "fp_incoming"),
        ("Production floor Area", ["production", "workshop", "fabrication", "floor"], "fp_production"),
        ("Packing Area", ["packing", "packaging area"], "fp_packing"),
        ("Inspection Area", ["inspection area", "qc area"], "fp_inspection"),
        ("Lab Area – Equipment used for Testing and inspection", ["lab", "laboratory", "testing equipment"], "fp_lab"),
        ("Calibration Labels on Equipment", ["calibration label", "cal label"], "fp_cal_label"),
        ("Non-Conformity Area / Rejection Area for RM & Final Product",
         ["rejection", "non-conform", "nonconform"], "fp_rejection"),
        ("Final Product Storage Area", ["final product storage", "finished storage", "fg storage"], "fp_final_storage"),
        ("Final Product Photos", ["final product", "product photo", "finished product photo",
                                    "product shape", "product identity", "package shape"], "fp_product_shape"),
        ("A clear image of the product label", ["product label", "label"], "fp_product_label"),
        ("A clear image of the product barcode / QR code", ["barcode", "qr code", "qr"], "fp_barcode"),
        ("An image showing identity and shape of the product / each model",
         ["product identity", "product shape", "model"], "fp_product_shape"),
        ("An image of the trademark", ["trademark", "brand logo", "logo"], "fp_trademark"),
    ]),
]


@dataclass
class MasterEntry:
    name: str
    doc_no: str
    rev: str = ""
    date: str = ""
    source: str = "master_list"


@dataclass
class ChecklistItemResult:
    section: str
    label: str
    present: bool
    matched_file: str = ""
    status: str = "not_found"


@dataclass
class ImageMatch:
    slot_id: str
    slot_label: str
    filename: str
    source_path: str


@dataclass
class DocumentMatch:
    heading: str
    filename: str
    doc_no: str = ""
    source: str = "folder"
    score: float = 0.0
    status: str = "provided_file"


@dataclass
class FAResult:
    folder: str
    master_list: str | None
    sections: dict[str, list[DocumentMatch]] = field(default_factory=dict)
    images: dict[str, list[ImageMatch]] = field(default_factory=dict)
    checklist: list[ChecklistItemResult] = field(default_factory=list)
    output_docx: str = ""
    summary_path: str = ""
    checklist_path: str = ""
    factory_info: dict[str, str] = field(default_factory=dict)
    audit_meta: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return {
            "folder": self.folder,
            "master_list": self.master_list,
            "sections": {
                sid: [
                    {
                        "heading": m.heading,
                        "filename": m.filename,
                        "doc_no": m.doc_no,
                        "source": m.source,
                        "status": m.status,
                    }
                    for m in matches
                ]
                for sid, matches in self.sections.items()
            },
            "images": {
                slot: [
                    {"slot_label": m.slot_label, "filename": m.filename, "source_path": m.source_path}
                    for m in matches
                ]
                for slot, matches in self.images.items()
            },
            "checklist": [
                {
                    "section": c.section,
                    "label": c.label,
                    "present": c.present,
                    "matched_file": c.matched_file,
                    "status": c.status,
                }
                for c in self.checklist
            ],
            "output_docx": self.output_docx,
            "summary_path": self.summary_path,
            "checklist_path": self.checklist_path,
            "factory_info": self.factory_info,
            "audit_meta": self.audit_meta,
        }


MASTER_LIST_ONLY_FILENAME = "(listed in master list - file not uploaded)"


def _is_master_list_only_filename(filename: str) -> bool:
    return filename in {
        MASTER_LIST_ONLY_FILENAME,
        "(on master list — file not in folder)",
    }


def find_master_list(folder: Path) -> Path | None:
    candidates: list[tuple[int, Path]] = []
    for path in folder.rglob("*"):
        if not path.is_file():
            continue
        if path.suffix.lower() not in {".pdf", ".docx"}:
            continue
        name = path.name.lower()
        score = 0
        for pat in MASTER_LIST_PATTERNS:
            if re.search(pat, name, re.I):
                score += 10
        if "procedure" in name and "list" in name:
            score += 8
        if score:
            candidates.append((score, path))
    if not candidates:
        return None
    candidates.sort(key=lambda x: (-x[0], len(x[1].name)))
    return candidates[0][1]


def parse_master_list_docx(docx_path: Path) -> list[MasterEntry]:
    """Parse a DOCX master list whose table columns are:
    Activity/Section | Required Documents | Documents Shared

    Each row maps a checklist section to the actual document filenames provided.
    Returns MasterEntry objects with the shared filename as doc_no for easy file lookup.
    """
    try:
        from docx import Document as _Doc
        doc = _Doc(str(docx_path))
    except Exception:
        return []

    entries: list[MasterEntry] = []
    section_kw_map = {
        "incoming material": ["2.1", "2.2", "2.3"],
        "process control":   ["3.2", "3.3", "3.4"],
        "calibration":       ["4.1"],
        "100% inspection":   ["5.2", "6.1"],
        "general":           [],   # handled separately
        "factory photos":    [],   # handled separately
        "final product photo": [],
    }

    for tbl in doc.tables:
        for ri, row in enumerate(tbl.rows):
            if ri == 0:
                continue   # header row
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 2:
                continue
            section_label = cells[0].lower()
            shared_raw = cells[2] if len(cells) > 2 else cells[1]
            # Shared col may have multiple filenames separated by newlines
            for shared in re.split(r"\n+", shared_raw):
                shared = shared.strip()
                if not shared or shared.lower() in ("n/a", "-", ""):
                    continue
                # Use the filename stem as a pseudo doc_no so match_documents can find the file
                doc_no = Path(shared).stem
                entries.append(MasterEntry(name=shared, doc_no=doc_no))
    return entries


FACTORY_INFO_FILE_PATTERNS = [
    r"factory\s*info",
    r"factory\s*information",
    r"factory\s*details",
    r"factory\s*audit.*application",
    r"application\s*form",
]

TEMPLATE_FACTORY_LABELS: dict[str, list[str]] = {
    "registered_name": ["Manufacturer’s registered name:", "Manufacturer's registered name:"],
    "street": ["Street and no.:"],
    "postal_code": ["Postal code:"],
    "city": ["City:"],
    "province": ["Province:"],
    "country": ["Country:"],
    "representative_name": ["Manufacturer's representative name:", "Manufacturer’s representative name:"],
    "position": ["Position:"],
    "telephone": ["Telephone:"],
    "fax": ["Fax:"],
    "email": ["E-Mail:", "E-mail:"],
    "factory_area": ["Factory Area Size:"],
    "employees": ["Number of employees:"],
    "products_manufactured": ["Products manufactured:"],
    "product_name": ["Product name:"],
    "product_brand": ["Product brand/trademark:"],
    "subcontractors": ["Main Subcontractors:"],
}


def find_factory_info_file(folder: Path) -> Path | None:
    candidates: list[tuple[int, Path]] = []
    for path in folder.rglob("*"):
        if not path.is_file():
            continue
        name = path.name.lower()
        if path.suffix.lower() not in {".pdf", ".docx"}:
            continue
        score = 0
        if re.search(r"factory\s*info|factory\s*information|factory\s*details", name, re.I):
            score += 30
        elif re.search(r"audit\s*plan|factory\s*audit|request\s*for\s*inspection", name, re.I):
            score += 20
        elif re.search(r"factory\s*audit.*application|application\s*form", name, re.I):
            score += 15
        if path.suffix.lower() == ".docx":
            score += 5   # prefer DOCX for table-based parsing
        if score:
            candidates.append((score, path))
    if not candidates:
        return None
    candidates.sort(key=lambda x: (-x[0], len(x[1].name)))
    return candidates[0][1]


def _pdf_text(path: Path, max_pages: int = 5) -> str:
    try:
        reader = PdfReader(str(path))
        return "\n".join((p.extract_text() or "") for p in reader.pages[:max_pages])
    except Exception:
        return ""


def _is_audit_plan_path(path: Path) -> bool:
    return bool(re.search(r"audit[\s_]*plan", path.name, re.I))


def _tesseract_pdf_text(path: Path, max_pages: int = 6, zoom: float = 4.0) -> str:
    """High-quality OCR for scanned audit plans (tesseract PSM 3 on rendered pages)."""
    if pytesseract is None:
        return ""
    try:
        import fitz

        doc = fitz.open(str(path))
        parts: list[str] = []
        mat = fitz.Matrix(zoom, zoom)
        for i in range(min(max_pages, doc.page_count)):
            page = doc.load_page(i)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = Image.open(BytesIO(pix.tobytes("png")))
            gray = ImageOps.autocontrast(ImageOps.grayscale(img))
            parts.append(pytesseract.image_to_string(gray, config="--psm 3"))
        return "\n".join(parts).strip()
    except Exception:
        return ""


def _audit_priority_docs(docs: list[tuple[Path, str]]) -> list[tuple[Path, str]]:
    audit_docs = [
        (path, text)
        for path, text in docs
        if re.search(r"audit[\s_]*plan|factory\s*audit|audit\s*application|request\s*for\s*inspection", path.name, re.I)
        or re.search(r"Factory Audit|Witness Testing|audit was carried out|KSA_Saleem", text, re.I)
    ]
    audit_docs.sort(
        key=lambda item: (
            0 if _is_audit_plan_path(item[0]) else 1,
            item[0].name.lower(),
        )
    )
    return audit_docs


def _docx_text(path: Path) -> str:
    try:
        with zipfile.ZipFile(path) as zin:
            xml = zin.read("word/document.xml")
        root = ET.fromstring(xml)
        return "\n".join(t.text or "" for t in root.findall(f".//{W}t"))
    except Exception:
        return ""


def _xlsx_text(path: Path) -> str:
    """Best-effort XLSX text extraction without adding a dependency."""
    try:
        with zipfile.ZipFile(path) as zin:
            parts: list[str] = []
            for name in zin.namelist():
                if name == "xl/sharedStrings.xml" or re.match(r"xl/worksheets/sheet\d+\.xml", name):
                    try:
                        root = ET.fromstring(zin.read(name))
                    except Exception:
                        continue
                    parts.extend(t.text or "" for t in root.iter() if t.tag.endswith("}t"))
            return "\n".join(parts)
    except Exception:
        return ""


def _document_text(path: Path, max_pages: int = 5) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        text = _pdf_text(path, max_pages=max_pages)
        direct_text = text
        if len(text.strip()) < 40:
            try:
                import fitz

                doc = fitz.open(str(path))
                text = "\n".join(doc.load_page(i).get_text() for i in range(min(max_pages, doc.page_count)))
                direct_text = text
                if len(text.strip()) < 80 and hasattr(fitz.Page, "get_textpage_ocr"):
                    ocr_texts: list[str] = []
                    for i in range(min(max_pages, doc.page_count)):
                        page = doc.load_page(i)
                        try:
                            tp = page.get_textpage_ocr(language="eng", dpi=200, full=True)
                            ocr_texts.append(page.get_text("text", textpage=tp))
                        except Exception:
                            continue
                    ocr_text = "\n".join(ocr_texts).strip()
                    if ocr_text:
                        text = ocr_text
            except Exception:
                pass
        # Keep typed dates/names from direct extraction when OCR loses them.
        preserved: list[str] = []
        for line in direct_text.splitlines():
            line = line.strip()
            if not line:
                continue
            if re.search(r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b", line):
                if line.lower() not in text.lower():
                    preserved.append(line)
            elif re.search(r"\b(?:auditor|audit member|lead auditor)\b", line, re.I):
                if line.lower() not in text.lower():
                    preserved.append(line)
        if preserved:
            text = "\n".join(preserved + [text])
        if _is_audit_plan_path(path) or (
            len(text.strip()) < 120
            and re.search(r"factory\s*audit|audit[\s_]*plan", path.name, re.I)
        ):
            tess = _tesseract_pdf_text(path, max_pages=max_pages)
            if len(tess.strip()) > len(text.strip()):
                text = "\n".join(preserved + [tess]) if preserved else tess
        return text
    if ext == ".docx":
        return _docx_text(path)
    if ext == ".xlsx":
        return _xlsx_text(path)
    return ""


def _iter_document_texts(folder: Path) -> list[tuple[Path, str]]:
    docs: list[tuple[Path, str]] = []
    for path in folder.rglob("*"):
        if not path.is_file() or path.name.startswith("."):
            continue
        if path.name.startswith("FA_"):
            continue
        if path.suffix.lower() not in {".pdf", ".docx", ".xlsx"}:
            continue
        text = _document_text(path, max_pages=6)
        if text.strip():
            docs.append((path, text))
    return docs


def _parse_factory_info_docx_tables(path: Path) -> dict[str, str]:
    """Read factory-info fields directly from the DOCX table cells.

    The form uses merged cells that repeat a value across columns, which breaks
    flattened-text regex parsing. Reading rows/cells directly is reliable.
    """
    label_map = [
        ("registered_name", "manufacturer’s registered name"),
        ("registered_name", "manufacturer's registered name"),
        ("street", "street and no"),
        ("postal_code", "postal code"),
        ("city", "city"),
        ("province", "province"),
        ("country", "country"),
        ("representative_name", "manufacturer’s representative name"),
        ("representative_name", "manufacturer's representative name"),
        ("position", "position"),
        ("telephone", "telephone"),
        ("email", "e-mail"),
        ("factory_area", "factory area size"),
        ("employees", "number of employees"),
        ("products_manufactured", "products manufactured"),
        ("product_name", "product name"),
        ("product_brand", "product brand/trademark"),
        ("subcontractors", "main subcontractors"),
    ]
    try:
        from docx import Document as _Doc
        doc = _Doc(str(path))
    except Exception:
        return {}

    info: dict[str, str] = {}

    def first_value(cells: list[str], skip_idx: int) -> str:
        for j, c in enumerate(cells):
            if j <= skip_idx:
                continue
            v = c.strip()
            if v and v.lower() != cells[skip_idx].strip().lower():
                return v
        return ""

    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            for idx, cell in enumerate(cells):
                cl = cell.strip().lower().rstrip(":")
                for key, label in label_map:
                    if key in info:
                        continue
                    if cl == label:
                        val = first_value(cells, idx)
                        if val:
                            info[key] = re.sub(r"\s+", " ", val).strip(" .;:")
                        break
    return _sanitize_factory_info(info)


def _clean_factory_info_value(value: str) -> str:
    value = re.sub(r"\s+", " ", value or "").strip(" .;:")
    value = re.sub(r"\b(?:click here to enter text|select|choose)\b.*$", "", value, flags=re.I).strip(" .;:")
    return value


def _valid_factory_name(value: str) -> bool:
    value = _clean_factory_info_value(value)
    if len(value) < 3 or len(value) > 120:
        return False
    if re.search(r"\b(e-?mail|telephone|phone|postal|province|country|representative|signature|auditor)\b", value, re.I):
        return False
    if re.search(r"\b(?:eee|ne eee)\b", value, re.I):
        return False
    if sum(ch.isdigit() for ch in value) > max(2, len(value) // 5):
        return False
    words = re.findall(r"[A-Za-z]{2,}", value)
    if not words:
        return False
    # PDF OCR garbage often has one mixed/camel fragment and no real company word.
    if len(words) <= 2 and not re.search(r"\b(?:ltd|limited|llc|co\.?|company|factory|industries|textile|manufacturing|pvt|inc)\b", value, re.I):
        return False
    return True


def _sanitize_factory_info(info: dict[str, str]) -> dict[str, str]:
    cleaned: dict[str, str] = {}
    for key, value in info.items():
        if key.startswith("_"):
            cleaned[key] = value
            continue
        value = _clean_factory_info_value(value)
        if not value:
            continue
        if key == "registered_name" and not _valid_factory_name(value):
            continue
        cleaned[key] = value
    return cleaned


def _parse_factory_info_pdf_fields(path: Path) -> dict[str, str]:
    """Read factory-info values from AcroForm fields when a PDF exposes them."""
    try:
        reader = PdfReader(str(path))
        fields = reader.get_fields() or {}
    except Exception:
        return {}
    if not fields:
        return {}

    key_patterns: list[tuple[str, str]] = [
        ("registered_name", r"registered.*name|manufacturer.*name|company.*name|factory.*name"),
        ("street", r"street|address|location"),
        ("postal_code", r"postal"),
        ("city", r"city"),
        ("province", r"province|state"),
        ("country", r"country"),
        ("representative_name", r"representative|contact\s*person"),
        ("position", r"position|designation|title"),
        ("telephone", r"telephone|phone|mobile"),
        ("fax", r"fax"),
        ("email", r"e-?mail|email"),
        ("factory_area", r"factory.*area|area.*size"),
        ("employees", r"employee"),
        ("products_manufactured", r"products?\s*manufactured|product\s*name"),
        ("product_brand", r"brand|trademark"),
        ("subcontractors", r"subcontract"),
    ]
    info: dict[str, str] = {}
    for raw_name, field in fields.items():
        label = re.sub(r"[_\-.]+", " ", str(raw_name)).lower()
        try:
            value = str(field.get("/V", "") or "").strip()
        except Exception:
            continue
        value = re.sub(r"\s+", " ", value).strip(" .;:")
        if not value or value.lower() in {"off", "no", "false"}:
            continue
        for key, pattern in key_patterns:
            if key not in info and re.search(pattern, label, re.I):
                info[key] = value
                break
    return _sanitize_factory_info(info)


def parse_factory_info(folder: Path) -> dict[str, str]:
    info_path = find_factory_info_file(folder)
    if not info_path:
        return {}

    # Prefer robust table-based parsing for DOCX forms
    if info_path.suffix.lower() == ".docx":
        table_info = _parse_factory_info_docx_tables(info_path)
        if table_info.get("registered_name") or table_info.get("country"):
            table_info["_source_file"] = info_path.name
            return table_info

    if info_path.suffix.lower() == ".pdf":
        field_info = _parse_factory_info_pdf_fields(info_path)
        if field_info.get("registered_name") or field_info.get("country"):
            field_info["_source_file"] = info_path.name
            return field_info

    if info_path.suffix.lower() == ".docx":
        text = _docx_text(info_path)
    else:
        text = _document_text(info_path, max_pages=6)
    if len(text.strip()) < 40:
        try:
            import fitz

            doc = fitz.open(str(info_path))
            text = "\n".join(doc.load_page(i).get_text() for i in range(min(3, doc.page_count)))
        except Exception:
            pass

    text = re.sub(r"\s+", " ", text)
    section = re.search(
        r"Manufacturer['\u2019]?s registered name and factory location",
        text,
        re.I,
    )
    if section:
        text = text[section.end():]
    ordered_fields = [
        ("registered_name", r"Manufacturer['\u2019]?s registered name:\s*"),
        ("street", r"Street and no\.:\s*"),
        ("postal_code", r"Postal code:\s*"),
        ("city", r"City:\s*"),
        ("province", r"Province:\s*"),
        ("country", r"Country:\s*"),
        ("_factory_rep_hdr", r"Factory representative name and contact data\s*"),
        ("representative_name", r"Manufacturer['\u2019]?s representative name:\s*"),
        ("position", r"Position:\s*"),
        ("telephone", r"Telephone:\s*"),
        ("fax", r"Fax:\s*"),
        ("email", r"E-?Mail:\s*"),
        ("_audit_team_hdr", r"Additional audit team member\(s\):\s*"),
        ("_factory_details_hdr", r"Factory details\s*"),
        ("factory_area", r"Factory Area Size:\s*"),
        ("employees", r"Number of employees:\s*"),
        ("products_manufactured", r"Products manufactured:\s*"),
        ("product_name", r"Product name:\s*"),
        ("product_brand", r"Product brand/trademark:\s*"),
        ("subcontractors", r"Main Subcontractors:\s*"),
    ]
    info: dict[str, str] = {}
    for i, (key, start_pat) in enumerate(ordered_fields):
        m = re.search(start_pat, text, re.I)
        if not m:
            continue
        if key.startswith("_"):
            continue
        start = m.end()
        end = len(text)
        if i + 1 < len(ordered_fields):
            nxt = re.search(ordered_fields[i + 1][1], text[start:], re.I)
            if nxt:
                end = start + nxt.start()
        value = text[start:end].strip(" .;:")
        if value:
            info[key] = value

    if not info:
        factory_section = re.search(r"FACTORY DETAILS\s*(.+?)(?:\s+\d+\s+PAYING PARTY|\s+\d+\s+Other Information|$)", text, re.I)
        section_text = factory_section.group(1) if factory_section else text
        name = _find_first([r"Company Name\s*[:\-]\s*([A-Z0-9][A-Za-z0-9 .,&/-]{2,90})"], section_text)
        rep = _find_first([r"Contact Person\s*[:\-]\s*([A-Z][A-Za-z .'-]{2,80})"], section_text)
        near = _find_first([r"Near\s*[:\-]\s*(.+?)(?:\s+Country\s*:|$)"], section_text)
        country = _find_first([r"Country\s*[:\-]\s*([A-Za-z ]{2,40})"], section_text)
        city = _find_first([r"City\s*[:\-]\s*([A-Za-z ]{2,40})"], section_text)
        products = re.findall(r"Click here to enter text\.\s*([A-Za-z0-9 /&().,-]{3,80})\s+Click here", text, re.I)
        if name:
            info["registered_name"] = name
        if rep:
            info["representative_name"] = rep
        if near:
            info["street"] = near
        if country:
            info["country"] = country.strip()
        if city:
            info["city"] = city.strip()
        if products:
            info["products_manufactured"] = ", ".join(dict.fromkeys(p.strip(" .") for p in products))
    if info:
        info["_source_file"] = info_path.name
    return _sanitize_factory_info(info)


COUNTRY_CODES: dict[str, str] = {
    "united arab emirates": "AE", "uae": "AE",
    "saudi arabia": "SA", "ksa": "SA",
    "hong kong": "HK",
    "china": "CN",
    "india": "IN",
    "united kingdom": "GB", "uk": "GB",
    "united states": "US", "usa": "US",
    "germany": "DE",
    "france": "FR",
    "italy": "IT",
    "spain": "ES",
    "netherlands": "NL",
    "belgium": "BE",
    "turkey": "TR",
    "egypt": "EG",
    "kuwait": "KW",
    "bahrain": "BH",
    "oman": "OM",
    "qatar": "QA",
    "jordan": "JO",
    "taiwan": "TW",
    "south korea": "KR",
    "japan": "JP",
    "malaysia": "MY",
    "singapore": "SG",
    "indonesia": "ID",
    "thailand": "TH",
    "vietnam": "VN",
    "pakistan": "PK",
    "bangladesh": "BD",
}


def _country_to_code(country: str) -> str:
    if not country:
        return ""
    c = country.strip().lower()
    if code := COUNTRY_CODES.get(c):
        return code
    # ISO-2 already given
    if re.match(r"^[A-Z]{2}$", country.strip()):
        return country.strip().upper()
    # try partial match
    for name, code in COUNTRY_CODES.items():
        if name in c or c in name:
            return code
    return country.strip().upper()[:2]


def _auditor_initials(auditor_names: list[str]) -> str:
    """Use only the lead auditor's initials (first name in the list)."""
    if not auditor_names:
        return ""
    parts = re.sub(r"[^A-Za-z ]", " ", auditor_names[0]).split()
    return "".join(p[0].upper() for p in parts if p)


def _normalize_auditor_code(value: str) -> str:
    return re.sub(r"[^A-Za-z]", "", value).upper()


def _expand_auditor_overrides(values: list[str] | None) -> tuple[list[str], list[str]]:
    """Return (full_names, codes) from portal-entered codes/names."""
    if not values:
        return [], []
    names: list[str] = []
    codes: list[str] = []
    for raw in values:
        raw = str(raw).strip()
        if not raw:
            continue
        code = _normalize_auditor_code(raw)
        if code in AUDITOR_REGISTRY:
            codes.append(code)
            names.append(AUDITOR_REGISTRY[code])
        else:
            names.append(re.sub(r"\s+", " ", raw))
    return names[:2], codes[:2]


def _auditor_signature_path(code: str) -> str:
    if not code:
        return ""
    for ext in ("png", "jpg", "jpeg", "webp"):
        path = AUDITOR_SIGNATURE_DIR / f"{code.upper()}.{ext}"
        if path.exists():
            return str(path)
    return ""


def generate_report_number(
    country: str,
    auditor_names: list[str],
    audit_date: str,
    seq: int = 1,
) -> str:
    """Format: FA + <country code> + <auditor initials> + DDMMYYYY + <seq 01>.
    Example: FAAENP2506202601

    Returns empty string if any required field (country, auditor initials, date)
    is missing — never generates a partial/assumed report number.
    """
    cc = _country_to_code(country)
    initials = _auditor_initials(auditor_names)
    date_compact = ""
    m = re.search(r"(\d{1,2})[/-](\d{1,2})[/-](\d{4})", audit_date or "")
    if m:
        d, mo, y = m.groups()
        date_compact = f"{int(d):02d}{int(mo):02d}{y}"

    # Do NOT produce a partial number — all three parts are required
    if not cc or not initials or not date_compact:
        return ""
    return f"FA{cc}{initials}{date_compact}{seq:02d}"


MONTHS = {
    "jan": "01", "january": "01",
    "feb": "02", "february": "02",
    "mar": "03", "march": "03",
    "apr": "04", "april": "04",
    "may": "05",
    "jun": "06", "june": "06",
    "jul": "07", "july": "07",
    "aug": "08", "august": "08",
    "sep": "09", "sept": "09", "september": "09",
    "oct": "10", "october": "10",
    "nov": "11", "november": "11",
    "dec": "12", "december": "12",
}


def _normalize_date(text: str) -> str:
    text = text.strip()
    m = re.search(r"\b(\d{1,2})[/-](\d{1,2})[/-](\d{2,4})\b", text)
    if m:
        d, mo, y = m.groups()
        if len(y) == 2:
            y = "20" + y
        return f"{int(d):02d}/{int(mo):02d}/{y}"
    m = re.search(r"\b(\d{1,2})(?:st|nd|rd|th)?\s+([A-Za-z]+)\s+(\d{4})\b", text, re.I)
    if m:
        d, mo, y = m.groups()
        mo_num = MONTHS.get(mo.lower())
        if mo_num:
            return f"{int(d):02d}/{mo_num}/{y}"
    return text


def _find_first(patterns: list[str], text: str) -> str:
    for pattern in patterns:
        m = re.search(pattern, text, re.I | re.S)
        if m:
            value = re.sub(r"\s+", " ", m.group(1)).strip(" :;,.")
            if value:
                return value
    return ""


def _normalize_person_name(value: str) -> str:
    value = re.sub(r"\s+", " ", value).strip(" .,/")
    parts = value.split()
    if len(parts) < 2:
        return ""
    return " ".join(part.capitalize() for part in parts)


def _clean_regulation_value(value: str) -> str:
    value = re.sub(r"\s+", " ", value).strip(" .;:")
    value = re.sub(r"^the\s+", "", value, flags=re.I)
    value = re.sub(r"^Technical Regulation for\s+", "", value, flags=re.I)
    if re.match(r"(?:the\s+)?(?:factory\s+will\s+be\s+visited|audit\s+was\s+carried|result\s+of\s+the\s+factory\s+audit)\b", value, re.I):
        return ""
    if re.search(r"\bSASO\s+Saleem\s+Program\b", value, re.I):
        return ""
    value = re.split(
        r"\s+(?:Lead Auditor|Audit Member|Observer|Product Scope|Factory Name|"
        r"Manufacturer|Audit Date|Auditor|Report No\.?|Date:)\b",
        value,
        maxsplit=1,
        flags=re.I,
    )[0].strip(" .;:")
    value = re.sub(
        r"\bRegulation\s*\(s\)\s+and\s+their\b",
        "and their",
        value,
        flags=re.I,
    )
    value = re.sub(r"\bprocedure\b.*$", "", value, flags=re.I).strip(" .;")
    return value


def _extract_technical_regulation_from_text(text: str) -> str:
    compact = re.sub(r"\s+", " ", text)
    patterns = [
        r"Technical\s+(?:Technical\s+)?Regulation\s+for\s+(.+?)(?:\s+Lead Auditor|\s+Audit Member|\s+Observer\b|\s+Product Scope\b|\s+Factory Name\b|$)",
        r"Technical Regulation Name\s*[:\-]?\s*(.+?)(?:\s+Lead Auditor|\s+Audit Member|\s+Product Scope\b|\s+Factory Name\b|$)",
        r"Technical Regulations?\s*[:\-]?\s*(.+?)(?:\s+Lead Auditor|\s+Product Scope\b|\s+Factory Name\b|$)",
        r"(Technical Regulation for [A-Za-z0-9][A-Za-z0-9 /&(),.'\-]{8,160})",
        r"in conformity to\s+(the\s+Saudi Technical Regulation)\b",
        r"\b(Saudi\s+Technical\s+Regulation(?:\s+for\s+[A-Za-z0-9 /&(),.'\-]{3,120})?)",
    ]
    for pattern in patterns:
        m = re.search(pattern, compact, re.I)
        if not m:
            continue
        value = _clean_regulation_value(m.group(1))
        if len(value) >= 10:
            return value
    return ""


def _find_audit_date(docs: list[tuple[Path, str]]) -> str:
    audit_docs = _audit_priority_docs(docs)
    dated: list[tuple[int, str]] = []
    for _, text in audit_docs:
        for m in re.finditer(
            r"\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b|\b(\d{1,2}(?:st|nd|rd|th)?\s+[A-Za-z]{3,9}\s+\d{4})\b",
            text,
            re.I,
        ):
            raw = m.group(1) or m.group(2)
            window = text[max(0, m.start() - 80) : m.end() + 40]
            priority = 0
            if re.search(r"\bdate\b|designation|representative|signature", window, re.I):
                priority = 0
            elif re.search(r"audit date", window, re.I):
                priority = 1
            else:
                priority = 2
            dated.append((priority, _normalize_date(raw)))
    if dated:
        dated.sort(key=lambda item: item[0])
        return dated[0][1]

    for _, text in audit_docs:
        compact = re.sub(r"\s+", " ", text)
        m = re.search(r"Audit Date\s+From\s+(\d{8})\s+to\s+(\d{8})", compact, re.I)
        if m:
            raw = m.group(1)
            return f"{raw[:2]}/{raw[2:4]}/{raw[4:]}"
    return ""


def _find_technical_regulation(docs: list[tuple[Path, str]]) -> str:
    audit_docs = _audit_priority_docs(docs)
    other_docs = [(path, text) for path, text in docs if (path, text) not in audit_docs]
    for _, text in audit_docs + other_docs:
        value = _extract_technical_regulation_from_text(text)
        if value:
            return value
    return ""


def _looks_like_person_name(value: str) -> bool:
    if not value or len(value.split()) < 2:
        return False
    if re.search(
        r"\b(?:technical|auditor|lead|auditear|teetwical|audiear|activity|section|observer|member)\b",
        value,
        re.I,
    ):
        return False
    return bool(re.match(r"^[A-Za-z][A-Za-z .'-]+$", value))


def _parse_auditor_name_list(value: str) -> list[str]:
    names: list[str] = []
    value = re.split(r"\b(?:Factory|Date|Signature|Position|Contact|Designation|Time)\b", value, maxsplit=1, flags=re.I)[0]
    for part in re.split(r"/|,|&|\band\b", value):
        part = re.sub(
            r"^.*?\b(?:technical|teetwical|auditor|auditear|lead)\b\s*",
            "",
            part,
            flags=re.I,
        )
        clean = _normalize_person_name(part)
        if _looks_like_person_name(clean) and clean.lower() not in {n.lower() for n in names}:
            names.append(clean)
    return names


def _extract_auditor_names_from_text(text: str) -> list[str]:
    names: list[str] = []
    compact = re.sub(r"\s+", " ", text)

    for pattern in [
        r"Name\s+[A-Z][A-Z .]+\s+Name\s+([A-Z][A-Z .]{3,40})",
    ]:
        for m in re.finditer(pattern, compact):
            clean = _normalize_person_name(m.group(1))
            if _looks_like_person_name(clean) and clean.lower() not in {n.lower() for n in names}:
                names.append(clean)
            if len(names) >= 2:
                return names[:2]

    lead = _find_first(
        [
            r"Lead Auditor[^A-Za-z0-9/]*(.*?)(?:\s+Audit Member|\s+Observer|\s+Activity/|$)",
            r"Auditor(?:'s)? name(?:s)?(?: \(printed letters\))?\s*[:\-]\s*([A-Z][A-Za-z .,&/-]{3,120})",
            r"T[ÜU]V Rheinland (?:Auditor|Representative)\s*[:\-]\s*([A-Z][A-Za-z .,&/-]{3,120})",
            r"Audit(?:or)? Team\s*[:\-]\s*([A-Z][A-Za-z .,&/-]{3,120})",
        r"Audit Member\(s\)\s*[:\-]?\s*([A-Z][A-Za-z .,&/-]{3,120})",
        r"Lead Auditor\s*[:\-]\s*([A-Z][A-Za-z .,&/-]{3,80})",
        ],
        compact,
    )
    if lead:
        for clean in _parse_auditor_name_list(lead):
            if clean.lower() not in {n.lower() for n in names}:
                names.append(clean)
            if len(names) >= 2:
                return names[:2]
    return names[:2]


def _find_auditor_names(docs: list[tuple[Path, str]]) -> list[str]:
    audit_docs = _audit_priority_docs(docs)
    other_docs = [(path, text) for path, text in docs if (path, text) not in audit_docs]
    for _, text in audit_docs + other_docs:
        names = _extract_auditor_names_from_text(text)
        if names:
            return names[:2]
    return []


def _signature_candidates(folder: Path) -> dict[str, list[str]]:
    out = {"auditor": [], "factory": []}
    for path in folder.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in IMAGE_EXTENSIONS:
            continue
        blob = path.stem.lower()
        if not re.search(r"\b(sign|signature|sig)\b", blob):
            continue
        if re.search(r"factory|client|representative|manufacturer", blob):
            out["factory"].append(str(path))
        else:
            out["auditor"].append(str(path))

    extracted_dir = folder / "_extracted_signatures"

    def _extract_docx_signatures(path: Path) -> None:
        if not re.search(
            r"audit\s*plan|factory\s*audit|application\s*form|signature"
            r"|factory\s*info|general\s*info|factory\s*information|general\s*information",
            path.name,
            re.I,
        ):
            return
        try:
            with zipfile.ZipFile(path) as zin:
                rel_root = ET.fromstring(zin.read("word/_rels/document.xml.rels"))
                rel_map = {
                    rel.attrib.get("Id", ""): rel.attrib.get("Target", "")
                    for rel in rel_root
                }
                root = ET.fromstring(zin.read("word/document.xml"))
                extracted_dir.mkdir(exist_ok=True)
                n = 1
                for tbl in root.findall(f".//{W}tbl"):
                    for tr in tbl.findall(f"{W}tr"):
                        row_text = _cell_text(tr).lower()
                        if "signature" not in row_text:
                            continue
                        cells = tr.findall(f"{W}tc")
                        for idx, tc in enumerate(cells):
                            blips = tc.findall(f".//{A}blip")
                            if not blips:
                                continue
                            cell_text = _cell_text(tc).lower()
                            kind = "auditor"
                            if re.search(r"factory|representative|manufacturer|client", cell_text):
                                kind = "factory"
                            elif (
                                re.search(r"factory|representative|manufacturer|client", row_text)
                                and idx >= max(1, len(cells) // 2)
                            ):
                                kind = "factory"
                            elif re.search(r"auditor|t[üu]v|tuv|rheinland", cell_text):
                                kind = "auditor"
                            for blip in blips:
                                rid = blip.attrib.get(f"{R}embed")
                                target = rel_map.get(rid or "", "")
                                if not target:
                                    continue
                                media_name = target if target.startswith("word/") else f"word/{target}"
                                try:
                                    raw = zin.read(media_name)
                                except KeyError:
                                    continue
                                ext = Path(media_name).suffix.lstrip(".") or "png"
                                out_path = extracted_dir / f"{kind}_signature_{path.stem}_{n}.{ext}"
                                out_path.write_bytes(raw)
                                out[kind].append(str(out_path))
                                n += 1
        except Exception:
            return

    for path in folder.rglob("*.docx"):
        if path.is_file():
            _extract_docx_signatures(path)

    def _classify_signature_image(page, xref: int) -> str:
        """Classify an embedded signature image using nearby page text."""
        try:
            rects = page.get_image_rects(xref)
            if not rects:
                return "auditor"
            rect = rects[0]
            cx = (rect.x0 + rect.x1) / 2
            cy = (rect.y0 + rect.y1) / 2
            nearby: list[tuple[float, str]] = []
            for block in page.get_text("blocks"):
                bx0, by0, bx1, by1, txt, *_ = block
                if not str(txt).strip():
                    continue
                bcx = (bx0 + bx1) / 2
                bcy = (by0 + by1) / 2
                # Prefer labels on the same row or directly above/below the signature.
                if abs(bcy - cy) <= 140 or (0 <= cy - bcy <= 180):
                    dist = abs(bcx - cx) + abs(bcy - cy)
                    nearby.append((dist, str(txt).lower()))
            blob = " ".join(txt for _, txt in sorted(nearby)[:8])
            if re.search(r"factory|representative|manufacturer|client", blob, re.I):
                return "factory"
            if re.search(r"auditor|t[üu]v|tuv|rheinland", blob, re.I):
                return "auditor"
            # In common two-column sign-off rows, the factory representative cell
            # is on the right and the auditor cell is on the left.
            # Stamps/seals (roughly square) near no auditor label → factory.
            try:
                pil_tmp = Image.open(BytesIO(doc.extract_image(xref).get("image", b"")))
                _w, _h = pil_tmp.size
                if 0.7 <= _w / max(_h, 1) <= 1.6:
                    return "factory"
            except Exception:
                pass
            return "factory" if cx > page.rect.width * 0.55 else "auditor"
        except Exception:
            return "auditor"

    for path in folder.rglob("*.pdf"):
        if not path.is_file():
            continue
        if not re.search(
            r"audit\s*plan|factory\s*audit|application\s*form|signature"
            r"|factory\s*info|general\s*info|factory\s*information|general\s*information",
            path.name,
            re.I,
        ):
            continue
        try:
            import fitz

            doc = fitz.open(str(path))
            extracted_dir.mkdir(exist_ok=True)
            n = 1
            for page in doc:
                for img in page.get_images(full=True):
                    data = doc.extract_image(img[0])
                    raw = data.get("image")
                    ext = data.get("ext", "png")
                    if not raw:
                        continue
                    try:
                        pil = Image.open(BytesIO(raw))
                    except Exception:
                        continue
                    w, h = pil.size
                    # Allow larger images to catch stamps (e.g. circular company seal).
                    # Reject very large full-page images and tiny icons.
                    if w > 900 or h > 600 or w < 60 or h < 25:
                        continue
                    aspect = w / max(h, 1)
                    # Accept wide signatures AND roughly-square stamps/seals (aspect 0.7–1.6).
                    # Only skip if it looks like a non-signature logo (very tall or tiny aspect).
                    if aspect < 0.7:
                        continue
                    kind = _classify_signature_image(page, img[0])
                    out_path = extracted_dir / f"{kind}_signature_{path.stem}_{n}.{ext}"
                    out_path.write_bytes(raw)
                    out[kind].append(str(out_path))
                    n += 1

                # Fallback: page is a pure raster scan (no selectable text, one large image).
                # Crop the bottom-right area where factory rep signatures typically appear.
                all_imgs_on_page = page.get_images(full=True)
                is_pure_scan = (
                    len(all_imgs_on_page) == 1
                    and not page.get_text("blocks")
                )
                if is_pure_scan and not out["factory"]:
                    pw, ph = page.rect.width, page.rect.height
                    # Signature/stamp is almost always in the lower 35% of the page.
                    # Use the full width — factory rep can be on either side.
                    clip = fitz.Rect(0, ph * 0.65, pw, ph)
                    pix = page.get_pixmap(clip=clip, dpi=150)
                    raw_crop = pix.tobytes("png")
                    try:
                        pil_crop = Image.open(BytesIO(raw_crop))
                        cw, ch = pil_crop.size
                        if cw >= 80 and ch >= 60:
                            out_path = extracted_dir / f"factory_signature_{path.stem}_{n}.png"
                            out_path.write_bytes(raw_crop)
                            out["factory"].append(str(out_path))
                            n += 1
                    except Exception:
                        pass
        except Exception:
            continue

    out["auditor"] = sorted(dict.fromkeys(out["auditor"]))[:2]
    out["factory"] = sorted(dict.fromkeys(out["factory"]))[:1]
    return out


def extract_audit_meta(
    folder: Path,
    factory_info: dict[str, str],
    auditor_names_override: list[str] | None = None,
) -> dict[str, Any]:
    """Extract audit metadata strictly from documents found in *folder*.

    IMPORTANT — no defaults / assumptions:
    - audit_date is left blank if not found in any document
    - auditor_names is left empty if not found and no override is provided
    - report_number is left blank if country, auditor initials, or date is missing
    Fields are only populated when they are actually present in the uploaded documents
    or explicitly provided via the portal UI (auditor_names_override / audit_date override).
    """
    docs = _iter_document_texts(folder)
    all_text = "\n".join(text for _, text in docs)
    compact = re.sub(r"\s+", " ", all_text)

    factory_name = factory_info.get("registered_name", "")
    # Do not guess factory names from broad PDF text. Bad extraction is worse
    # than a blank field; use labelled factory-info fields only.

    signatures = _signature_candidates(folder)
    audit_date = _find_audit_date(docs)
    auditor_names = _find_auditor_names(docs)
    auditor_codes: list[str] = []
    if auditor_names_override:
        override_names, auditor_codes = _expand_auditor_overrides(auditor_names_override)
        if override_names:
            auditor_names = override_names[:2]

    code_signature_paths = [
        sig for sig in (_auditor_signature_path(code) for code in auditor_codes) if sig
    ]
    if code_signature_paths:
        signatures["auditor"] = code_signature_paths[:2]
    cover_code = next((code for code in auditor_codes if code in MAIN_AUDITOR_CODES), "")
    cover_signature = _auditor_signature_path(cover_code) if cover_code else (code_signature_paths[0] if code_signature_paths else "")
    country = factory_info.get("country", "")
    report_number = generate_report_number(country, auditor_names, audit_date)

    meta: dict[str, Any] = {
        "factory_name": factory_name,
        "technical_regulation": _find_technical_regulation(docs),
        "audit_date": audit_date,
        "auditor_names": auditor_names,
        "country": country,
        "factory_representative": factory_info.get("representative_name", ""),
        "auditor_signatures": signatures["auditor"],
        "cover_auditor_signature": cover_signature,
        "auditor_codes": auditor_codes,
        "factory_signature": signatures["factory"][0] if signatures["factory"] else "",
        "support_docs": detect_support_documents(folder),
        "report_number": report_number,
    }
    return {k: v for k, v in meta.items() if v}


def detect_support_documents(folder: Path) -> dict[str, Any]:
    files = list_folder_files(folder)
    file_blobs = [(p, _text_blob(p.stem, p.name.lower())) for p in files]

    def match_any(patterns: list[str]) -> str:
        for path, blob in file_blobs:
            if any(re.search(pattern, blob, re.I) for pattern in patterns):
                return str(path.resolve())
        return ""

    iso9001 = match_any([r"iso\s*9001", r"\b9k\b", r"9001.*cert", r"quality.*management.*cert"])
    iso14001 = match_any([r"iso\s*14001", r"\b14k\b", r"14001.*cert", r"environment.*management.*cert"])
    iso45001 = match_any([r"iso\s*45001", r"\b45k\b", r"45001.*cert", r"occupational.*health.*cert"])
    gmp = match_any([r"\bgmp\b", r"good manufacturing practice"])

    iso9001_details = _extract_iso_certificate_details(Path(iso9001), "9001") if iso9001 else {}

    return {
        "iso9001": bool(iso9001),
        "iso14001": bool(iso14001),
        "iso45001": bool(iso45001),
        "gmp": bool(gmp),
        "qms_certified": bool(iso9001),
        "iso9001_path": iso9001,
        "iso14001_path": iso14001,
        "iso45001_path": iso45001,
        "gmp_path": gmp,
        "iso9001_details": iso9001_details,
    }


def _extract_iso_certificate_details(path: Path, iso_no: str) -> dict[str, str]:
    """Best-effort extraction of ISO certificate details for section 8.0."""
    if not path.exists():
        return {}
    text = _document_text(path, max_pages=3)
    compact = re.sub(r"\s+", " ", text)
    standard = _find_first([rf"\b(ISO\s*{iso_no}\s*:\s*\d{{4}})\b"], compact)
    if not standard and re.search(rf"\bISO\s*{iso_no}\b", compact, re.I):
        standard = f"ISO {iso_no}"
    cert_no = _find_first(
        [
            r"(?:certificate|cert\.?)\s*(?:no|number|#)\s*[:\-]?\s*([A-Z0-9][A-Z0-9./\-]{3,40})",
            r"\bCertificate\s+([A-Z]{1,6}\d[\w./\-]{3,40})",
        ],
        compact,
    )
    body = _find_first(
        [
            r"(?:certification body|certified by|issued by)\s*[:\-]?\s*([A-Z][A-Za-z0-9 .,&/-]{3,80})",
            r"\b(Bureau Veritas Certification|Intertek|SGS|TUV Rheinland|TÜV Rheinland|DQS|BSI|LRQA)\b",
        ],
        compact,
    )
    dates = ""
    date_hits = re.findall(
        r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\d{1,2}\s+[A-Za-z]{3,9}\s+\d{4}\b",
        compact,
        re.I,
    )
    normalized_dates = [_normalize_date(d) for d in date_hits]
    if normalized_dates:
        unique_dates = list(dict.fromkeys(normalized_dates))
        if len(unique_dates) >= 2:
            dates = f"{unique_dates[0]} To {unique_dates[-1]}"
        else:
            dates = unique_dates[0]
    return {
        "standard": standard or f"ISO {iso_no}",
        "body": body or "NA",
        "cert_no": cert_no or "NA",
        "dates": dates or "NA",
    }


def parse_master_list(pdf_path: Path) -> list[MasterEntry]:
    """Parse master-list files (PDF or DOCX)."""
    if pdf_path.suffix.lower() == ".docx":
        return parse_master_list_docx(pdf_path)
    reader = PdfReader(str(pdf_path))
    text = "\n".join((p.extract_text() or "") for p in reader.pages)
    entries: list[MasterEntry] = []
    seen: set[str] = set()

    def _add(name: str, doc_no: str, rev: str = "", date: str = "") -> None:
        doc_no = doc_no.strip().upper()
        name = re.sub(r"\s+", " ", name).strip(" -–/")
        if not doc_no or not name or doc_no in seen:
            return
        # Skip header/legend noise
        if re.search(
            r"daftar\s+(kebijakan|dokumen)|revisi|tgl\b|doc\.?\s*no|"
            r"^report\s+name$|^documents?\s+control$|^process\s+area$|"
            r"^book\s*/?\s*computer$|^version$|^sl$",
            name,
            re.I,
        ):
            return
        if len(name) < 3 or not re.search(r"[A-Za-z]{3}", name):
            return
        seen.add(doc_no)
        entries.append(MasterEntry(name=name, doc_no=doc_no, rev=rev, date=date))

    # --- Format A: IMS English style  XXX-YYY-PR-001 ---
    for line in text.splitlines():
        line = re.sub(r"\s+", " ", line).strip()
        if not line or line.startswith("Sr #") or line.startswith("Level "):
            continue
        if line.startswith("Page ") or line.startswith("List of IMS"):
            continue
        m = re.search(
            r"^(.+?)\s+((?:[A-Z]{2,5}(?:-[A-Z]{2,5}){1,4}-(?:PR|FR|FC|ML|PY)-\d{3,4}))\s+(\S+)\s+([\d-]+(?:\s+.*)?)$",
            line,
            re.I,
        )
        if m:
            _add(m.group(1), m.group(2), m.group(3), m.group(4).strip())
            continue
        m2 = re.search(
            r"^(.+?)\s+((?:[A-Z]{2,5}(?:-[A-Z]{2,5}){1,4}-(?:PR|FR|FC|ML|PY)-\d{3,4}))",
            line,
            re.I,
        )
        if m2:
            _add(m2.group(1), m2.group(2))

    # Process maps on their own lines
    for line in text.splitlines():
        line = re.sub(r"\s+", " ", line).strip()
        pm = re.search(
            r"Process Map:\s*(.+?)\s*\(((?:[A-Z]{2,5}(?:-[A-Z]{2,5}){1,4}-(?:PR|FR|FC)-\d{3,4})\s*[^)]*)\)",
            line,
            re.I,
        )
        if pm:
            doc = DOC_NO_RE.search(pm.group(2))
            if doc:
                _add(pm.group(1).strip(), doc.group(1))

    # --- Format B: Wintai / Indonesian  WG/PRO-PROD/076  or  SMQ-WG-02 ---
    # Collapse soft line-breaks so wrapped titles become one line
    flat = re.sub(r"\s+", " ", text)
    # Pattern: optional row# + DOC.NO + TITLE + optional REV + DATE
    wg_pat = re.compile(
        r"(?:(?<=\s)|^)(?:\d{1,3}\s+)?"
        r"(WG/[A-Z]{2,6}-[A-Z]{2,6}/\d{3}|S[MLQ]{2}-WG-\d{2})"
        r"\s+(.+?)"
        r"(?=\s+\d{1,2}\s+\d{2}[-/]\d{2}[-/]\d{4}"   # rev + date
        r"|\s+\d{1,3}\s+WG/"                           # next row#
        r"|\s+\d{1,3}\s+S[MLQ]{2}-WG-"
        r"|DAFTAR\s+"
        r"|WG\s*:\s*Wintai"
        r"|$)",
        re.I,
    )
    for m in wg_pat.finditer(flat):
        doc_no = m.group(1)
        rest = m.group(2).strip()
        # Strip trailing revision number if present (single digit before date was consumed by lookahead)
        rest = re.sub(r"\s+\d{1,2}$", "", rest).strip()
        _add(rest, doc_no)

    # --- Format C: F-XXX style  (e.g. Belamy / Bangladesh factories) ---
    # Line format: NUM  REPORT_NAME  F-XXXX  VERSION-XX  BOOK/COMPUTER...
    fxxx_pat = re.compile(
        r"^\d{1,3}\s+(.+?)\s+(F-\d{3,4}(?:\.\d+)?)\s+VERSION-\d{2}",
        re.I | re.MULTILINE,
    )
    for m in fxxx_pat.finditer(text):
        name = re.sub(r"\s+", " ", m.group(1)).strip()
        doc_no = m.group(2).upper()
        _add(name, doc_no)

    # --- Format D: GENERIC fallback (works on any tabular master list) ---
    # Auto-detects a document/reference code per row without hardcoding the
    # company's specific scheme. Runs last, so it only fills rows the specific
    # formats above missed. Keep this permissive but guarded so prose lines
    # (addresses, headings) are not mistaken for document rows.
    for raw_line in text.splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip()
        if len(line) < 5:
            continue
        _parse_generic_master_row(line, _add)

    return entries


# A document/reference code: starts with a letter, then either uses separators
# (F-105, SOP-QC-001, WG/PRO-PROD/076, F=08, MS-0042909) or is letters+digits
# with no separator (FQAR001). Case-insensitive.
_GENERIC_CODE_RE = re.compile(
    r"(?<![A-Za-z0-9])("
    r"(?=[A-Za-z0-9=/._-]*\d)"                         # document codes must contain digits
    r"[A-Za-z][A-Za-z0-9]*(?:[=/._-][A-Za-z0-9]+)+"   # with separators
    r"|[A-Za-z]{1,6}\d{2,6}"                            # letters then digits, no sep
    r")(?![A-Za-z0-9/=._-]*[A-Za-z]{4})",              # avoid swallowing whole words
    re.I,
)

# Alpha prefixes that are NOT document codes (version/revision/date/layout columns).
_GENERIC_CODE_STOP = {
    "version", "revision", "rev", "page", "iso", "no", "sl", "id", "book",
    "computer", "card", "swatch", "print", "date", "tgl", "doc", "ver",
}

# Row must look like a data row, not prose: begins with a row number, OR the
# code is accompanied by a version/control/date marker somewhere on the line.
_GENERIC_ROWNUM_RE = re.compile(r"^\d{1,4}[\.\)]?\s+")
_GENERIC_CONTROL_MARKER_RE = re.compile(
    r"version|revisi|\brev\b|\bver\b|book|computer|card|swatch|print|"
    r"\d{2}[-/]\d{2}[-/]\d{2,4}|\bno[.:]?\s*\d",
    re.I,
)

_GENERIC_NAME_STOP_RE = re.compile(
    r"\s+(?:version|revisi|\brev\b|\bver\b|book|computer|card|swatch|print|"
    r"\d{1,2}[-/]\d{1,2}[-/]\d{2,4}).*$",
    re.I,
)

_GENERIC_HEADER_RE = re.compile(
    r"^(?:sl|s/?n|no|id|doc(?:ument)?\s*(?:no|id)?|report\s+name|"
    r"documents?\s+control|version|book|computer|process\s+area)(?:\s|$)",
    re.I,
)


def _parse_generic_master_row(line: str, add) -> None:
    """Extract (name, doc_no) from a single tabular master-list row, generically.

    Safe by design: only accepts rows that clearly look like document-register
    entries (row number prefix or a version/control/date marker present).
    """
    has_rownum = bool(_GENERIC_ROWNUM_RE.match(line))
    has_marker = bool(_GENERIC_CONTROL_MARKER_RE.search(line))
    if not (has_rownum or has_marker):
        return

    body = _GENERIC_ROWNUM_RE.sub("", line, count=1) if has_rownum else line
    if _GENERIC_HEADER_RE.match(body):
        return

    chosen = None
    for m in _GENERIC_CODE_RE.finditer(body):
        token = m.group(1)
        prefix = re.split(r"[=/._-]", token, maxsplit=1)[0].lower()
        if prefix in _GENERIC_CODE_STOP:
            continue
        # Reject pure version/date-like tokens (e.g. 01-2025) — need a letter start.
        if not re.match(r"^[A-Za-z]", token):
            continue
        chosen = m
        break
    if not chosen:
        return

    doc_no = chosen.group(1).upper()
    before = body[: chosen.start()].strip(" -–/·.")
    after = body[chosen.end() :].strip(" -–/·.")
    # Support both common table shapes:
    #   Title first:  "Incoming Inspection Report F-105 VERSION-01"
    #   Code first:   "SOP-QC-001 Incoming Inspection Procedure Rev 01"
    name = before or after
    name = _GENERIC_NAME_STOP_RE.sub("", name).strip(" -–/·.")
    # A real document name has words; skip if it's basically empty or numeric.
    if len(name) < 3 or not re.search(r"[A-Za-z]{3}", name):
        return
    if _GENERIC_HEADER_RE.match(name):
        return
    add(name, doc_no)


# Indonesian / English synonym terms used when scoring master-list titles against section keywords
_ID_EN_SYNONYMS: list[tuple[str, str]] = [
    ("prosedur", "procedure"),
    ("pengendalian", "control"),
    ("produk tidak sesuai", "non conforming product"),
    ("tidak sesuai", "non conforming"),
    ("penolakan", "rejection"),
    ("reject", "rejection"),
    ("pelatihan", "training"),
    ("kalibrasi", "calibration certificate"),
    ("monitoring alat ukur", "calibration certificate"),
    ("sertifikat kalibrasi", "calibration certificate"),
    ("keluhan", "complaint"),
    ("pengaduan", "complaint"),
    ("customer claim", "customer complaint"),
    ("claim", "complaint"),
    ("recall", "recall"),
    ("traceability", "traceability"),
    ("ketertelusuran", "traceability"),
    ("penanganan", "handling"),
    ("bahan baku", "raw material"),
    ("benang", "yarn"),
    ("gudang", "warehouse"),
    ("penyimpanan", "storage"),
    ("packing", "packing"),
    ("finishing", "packing"),
    ("inspeksi", "inspection"),
    ("pemeriksaan", "inspection"),
    ("pengujian", "testing"),
    ("test produk", "final test"),
    ("instruksi kerja", "work instruction"),
    ("ik :", "work instruction"),
    ("perawatan", "maintenance"),
    ("pemeliharaan", "maintenance"),
    ("ppm", "preventive maintenance"),
    ("tindakan perbaikan", "corrective action"),
    ("pencegahan", "corrective action"),
    ("incoming", "incoming"),
    ("penerimaan", "receiving"),
    ("supplier", "supplier"),
    ("mutu", "quality"),
    ("quality", "quality"),
    ("qc in line", "in line inspection"),
    ("qc end line", "final inspection"),
    ("perencanaan test produk", "final product test"),
    ("test produk", "final test"),
    ("final inspection", "final inspection"),
    ("production record", "production record"),
    ("record", "record"),
    ("penyimpanan", "storage"),
    ("benang", "yarn storage"),
    ("penerimaan barang", "incoming material inspection"),
    ("penerimaan dan penyimpanan benang", "raw material handling incoming"),
    ("penerimaan dan penyimpanan", "raw material handling"),
    ("benang habis", "yarn"),
    ("pengiriman", "shipping"),
    ("pengepakan", "packing"),
]


def _expand_master_blob(name: str, doc_no: str = "") -> str:
    """Blob for scoring: original title + English synonyms of Indonesian terms."""
    base = _text_blob(name, doc_no)
    extras: list[str] = []
    for id_term, en_term in _ID_EN_SYNONYMS:
        if id_term in base:
            extras.append(en_term)
    return _text_blob(base, *extras)


# Dept codes in WG doc numbers and which FA section IDs they are most relevant to
_DEPT_SECTION_AFFINITY: dict[str, list[str]] = {
    "PROD": ["2.1", "2.2", "2.3", "3.2", "3.3", "3.4", "3.5", "3.6",
             "4.1", "4.2", "4.3", "4.4", "5.1", "5.2", "6.1", "6.2", "6.3"],
    "HR":   ["3.1"],               # Training is HR territory
    "HS":   ["4.5"],               # Corrective/health-safety
    "ENV":  ["4.5"],               # Corrective/env
    "IK":   ["3.3", "3.4", "3.2", "5.2", "6.1"],  # Work instructions → production
}


def _dept_affinity_boost(doc_no: str, sid: str) -> float:
    """Small bonus when the document department matches the expected section."""
    # Extract dept code: WG/PRO-PROD/076 → PROD; WG/IK-PROD/035 → IK+PROD
    parts = re.split(r"[/\-]", doc_no.upper())
    bonus = 0.0
    for dept, sids in _DEPT_SECTION_AFFINITY.items():
        if dept in parts:
            if sid in sids:
                bonus += 3.0
            # Penalise HR appearing in incoming/production quality sections
            if dept == "HR" and sid in ("2.1", "2.2", "2.3", "3.2", "3.3", "3.5",
                                         "5.1", "5.2", "6.1", "6.2", "6.3"):
                bonus -= 4.0
    return bonus


def list_folder_files(folder: Path) -> list[Path]:
    skip = {".ds_store", "thumbs.db"}
    files: list[Path] = []
    for path in folder.rglob("*"):
        if not path.is_file():
            continue
        if path.name.startswith(".") or path.name.lower() in skip:
            continue
        if path.suffix.lower() in {".pdf", ".xlsx", ".xls", ".docx", ".doc", ".jpeg", ".jpg", ".png"}:
            files.append(path)
    return sorted(files, key=lambda p: p.name.lower())


def _text_blob(*parts: str) -> str:
    return " ".join(parts).lower()


def score_match(section_id: str, blob: str) -> float:
    score = 0.0
    for kw in SECTION_KEYWORDS.get(section_id, []):
        if kw in blob:
            score += len(kw.split()) * 2
    for hint in FILENAME_HINTS.get(section_id, []):
        if hint in blob:
            score += 3
    return score


def extract_doc_no_from_filename(name: str) -> str:
    m = DOC_NO_RE.search(name)
    if m:
        return m.group(1).upper()
    # Try leading doc-number prefix in filename
    stem = Path(name).stem
    mp = _FILENAME_DOC_NO_RE.match(stem)
    if mp:
        return mp.group(1).upper()
    return ""


# Patterns for doc numbers embedded in document text (header, footer, first/last page)
_DOC_NO_TEXT_PATTERNS = [
    r"(?:document\s*(?:no|number|ref)[.:]\s*)([A-Z0-9][\w.\-/]{2,30})",
    r"(?:doc(?:ument)?\s*#\s*)([A-Z0-9][\w.\-/]{2,30})",
    r"(?:ref(?:erence)?\s*(?:no|number)?[.:]\s*)([A-Z0-9][\w.\-/]{2,30})",
    r"(?:procedure\s*no[.:]\s*)([A-Z0-9][\w.\-/]{2,30})",
    r"(?:form\s*no[.:]\s*)([A-Z0-9][\w.\-/]{2,30})",
    r"\b([A-Z]{2,6}-[A-Z]{0,6}\d{2,6}(?:-\d{1,3})?)\b",   # e.g. SOP-QC-001, FR-014
]


def _extract_doc_no_from_content(path: Path) -> str:
    """Try to read a doc-number from the first or last page of the document."""
    try:
        ext = path.suffix.lower()
        if ext == ".pdf":
            reader = PdfReader(str(path))
            pages = reader.pages
            # Check first and last page
            texts = []
            if pages:
                texts.append(pages[0].extract_text() or "")
            if len(pages) > 1:
                texts.append(pages[-1].extract_text() or "")
            text = " ".join(texts)
        elif ext == ".docx":
            from docx import Document as _Doc
            doc = _Doc(str(path))
            # First and last paragraph
            paras = [p.text for p in doc.paragraphs if p.text.strip()]
            text = " ".join(paras[:6] + paras[-3:])
        elif ext in {".xlsx", ".xls", ".doc"}:
            return ""
        else:
            return ""

        text = re.sub(r"\s+", " ", text)
        for pat in _DOC_NO_TEXT_PATTERNS:
            m = re.search(pat, text, re.I)
            if m:
                candidate = m.group(1).strip().upper()
                # Reject pure dates or very short hits
                if not re.match(r"^\d{4,}$", candidate) and len(candidate) >= 4:
                    return candidate
    except Exception:
        pass
    return ""


# Noise suffixes commonly found in factory document filenames (versions, audit codes)
_VERSION_NOISE_RE = re.compile(
    r"\s+(?:vs?\s*\d[\d.]*|v\s*\d[\d.]*|SO\s+\d[\d\-–/]+|rev[\s._]\w+|\(\d{4}\))\b.*$",
    re.I,
)


def _clean_doc_heading(heading: str, filename: str = "") -> str:
    """Return a short, clean document name (strip version/date noise, version codes)."""
    # Prefer stem from filename if heading is just the stem
    stem = Path(filename).stem if filename else ""
    raw = heading or stem
    # Strip the leading doc-number if present
    raw = _FILENAME_DOC_NO_RE.sub("", raw).strip()
    # Strip trailing version noise
    raw = _VERSION_NOISE_RE.sub("", raw).strip()
    # Remove trailing dash/underscore
    raw = re.sub(r"[\s\-_]+$", "", raw)
    return raw[:60] if raw else (stem[:60] if stem else heading[:60])


def match_documents(folder: Path) -> FAResult:
    folder = folder.resolve()
    master_path = find_master_list(folder)
    master_entries: list[MasterEntry] = []
    if master_path:
        try:
            master_entries = parse_master_list(master_path)
        except Exception:
            master_entries = []

    files = list_folder_files(folder)
    if master_path:
        files = [f for f in files if f.resolve() != master_path.resolve()]

    section_matches: dict[str, list[DocumentMatch]] = {sid: [] for sid in SECTION_IDS}
    used: dict[str, set[str]] = {sid: set() for sid in SECTION_IDS}

    def add_match(section_id: str, heading: str, filename: str, doc_no: str, source: str, score: float):
        key = f"{doc_no}|{filename}".lower()
        if key in used[section_id]:
            return
        # Also dedup by bare filename so master-list + direct-file don't both appear for same file
        fname_key = f"__fname__|{filename}".lower()
        is_master_only = _is_master_list_only_filename(filename)
        if filename and not is_master_only and fname_key in used[section_id]:
            return
        if score < 2 and not doc_no:
            return
        used[section_id].add(key)
        if filename and not is_master_only:
            used[section_id].add(fname_key)
        status = "master_list_only" if is_master_only else "provided_file"
        section_matches[section_id].append(
            DocumentMatch(heading=heading, filename=filename, doc_no=doc_no, source=source, score=score, status=status)
        )

    # Match master list entries to sections (primary source for 2.0–6.3 doc numbers)
    for entry in master_entries:
        blob = _expand_master_blob(entry.name, entry.doc_no)
        ranked: list[tuple[str, float]] = []
        for sid in SECTION_IDS:
            s = score_match(sid, blob) + _dept_affinity_boost(entry.doc_no, sid)
            if s > 0:
                ranked.append((sid, s))
        ranked.sort(key=lambda x: -x[1])
        if not ranked or ranked[0][1] < 4:
            continue
        top = ranked[0][1]
        # Link to an actual file in the folder when possible (by doc-no or title tokens)
        fname = ""
        _ML_STOP = {
            "prosedur", "procedure", "sop", "untuk", "dengan", "dalam", "yang",
            "test", "report", "record", "daily", "form", "check", "list",
            "sheet", "plan", "table", "data", "book", "card",
        }
        name_tokens = [
            t for t in re.split(r"[^a-z0-9]+", entry.name.lower())
            if len(t) >= 4 and t not in _ML_STOP
        ]
        for f in files:
            fl = f.name.lower()
            if _doc_no_matches_filename(entry.doc_no, fl):
                fname = f.name
                break
            if name_tokens:
                matched = sum(1 for t in name_tokens if t in fl)
                # Require majority of meaningful tokens to match (stricter than before)
                needed = max(2, len(name_tokens) * 2 // 3)
                if matched >= needed:
                    fname = f.name
                    break
        # Assign to best section; also allow NC sections to share if they score close
        for sid, s in ranked:
            if s < max(4, top * 0.75):
                continue
            # Only multi-assign for non-conformity sections; others take the single best
            if sid not in NON_CONFORMITY_SIDS and sid != ranked[0][0]:
                continue
            add_match(
                sid,
                entry.name,
                fname or MASTER_LIST_ONLY_FILENAME,
                entry.doc_no,
                "master_list",
                s + (5 if fname else 0),
            )
            if sid not in NON_CONFORMITY_SIDS:
                break  # one section only for normal docs

    # Files that belong to GENERAL DOCUMENTS (section 1) — must not leak into 2.0–6.3
    general_doc_re = re.compile(
        r"master\s*list|list of documents|organi[sz]ation\s*chart|org\s*chart|"
        r"organi[sz]ation\s*structure|business\s*license|trade\s*license|"
        r"manufacturing\s*license|\bnib\b|factory\s*layout|quality\s*manual|"
        r"iso\s*9001|iso\s*14001|iso\s*45001|certificate iso",
        re.I,
    )

    # Match folder files directly (skip raw image files — they go to photo slots only)
    for f in files:
        if f.suffix.lower() in IMAGE_EXTENSIONS:
            continue
        if general_doc_re.search(f.name):
            continue
        name_lower = f.name.lower()
        doc_no = extract_doc_no_from_filename(f.name)
        stem_blob = _path_match_blob(folder, f)
        if doc_no:
            stem_blob = _text_blob(stem_blob, doc_no)

        ranked: list[tuple[str, float]] = []
        for sid in SECTION_IDS:
            s = score_match(sid, stem_blob)
            if s > 0:
                ranked.append((sid, s))
        ranked.sort(key=lambda x: -x[1])

        if not ranked:
            continue

        top_score = ranked[0][1]
        for sid, s in ranked:
            if s < max(3, top_score * 0.6):
                continue
            heading = f.stem
            for entry in master_entries:
                if _doc_no_matches_filename(entry.doc_no, name_lower):
                    heading = entry.name
                    doc_no = doc_no or entry.doc_no
                    break
            add_match(sid, heading, f.name, doc_no, "folder", s)

    _SEC_MAX: dict[str, int] = {"4.1": 5, "3.4": 4, "3.1": 3, "3.6": 4, "4.3": 5}
    for sid in SECTION_IDS:
        section_matches[sid].sort(
            key=lambda m: (
                1 if m.status == "master_list_only" else 0,
                -m.score,
                m.heading.lower(),
            )
        )
        section_matches[sid] = section_matches[sid][:_SEC_MAX.get(sid, 2)]

    return FAResult(
        folder=str(folder),
        master_list=str(master_path) if master_path else None,
        sections=section_matches,
    )


def _score_keywords(blob: str, keywords: list[str]) -> float:
    score = 0.0
    for kw in keywords:
        if kw in blob:
            score += max(2, len(kw.split()) * 2)
    return score


def _doc_no_matches_filename(doc_no: str, filename: str) -> bool:
    """Match document codes to filenames without letting short fragments leak.

    Full document-code matches are always OK. Suffix-only matches are useful for
    filenames that omit a prefix, but only when the suffix is long enough to be
    distinctive (e.g. 1031, not 20 or 7).
    """
    if not doc_no or not filename:
        return False
    doc = doc_no.lower()
    fn = filename.lower()
    if doc in fn:
        return True
    suffix = re.split(r"[=/._-]", doc)[-1]
    return len(suffix) >= 3 and re.search(rf"(?<!\d){re.escape(suffix)}(?!\d)", fn) is not None


def _path_match_blob(folder: Path, path: Path) -> str:
    """Filename + parent folder path — so Docs/6. Factory Photos/1. Front Main Building/1.jpg matches."""
    try:
        rel = path.relative_to(folder.resolve())
        parts = [p for p in rel.parts[:-1] if p not in {".", "..", "Docs", "docs"}]
        return _text_blob(path.stem, path.name.lower(), " ".join(parts))
    except ValueError:
        return _text_blob(path.stem, path.name.lower(), path.parent.name)


def _find_checklist_file(
    keywords: list[str],
    files: list[Path],
    master_entries: list[MasterEntry],
    folder: Path | None = None,
) -> str:
    best_name = ""
    best_score = 0.0
    # Image files are OK when the match is driven by folder/name keywords
    # (e.g. Factory Layout.jpg, quality plan photos, calibration labels).
    allow_image_kw = any(
        k in " ".join(keywords).lower()
        for k in ("layout", "label", "tag", "photo", "plan", "chart", "record")
    )
    root = folder.resolve() if folder else None
    for f in files:
        if f.suffix.lower() in IMAGE_EXTENSIONS and not allow_image_kw:
            continue
        blob = _path_match_blob(root, f) if root else _text_blob(f.stem, f.name.lower())
        s = _score_keywords(blob, keywords)
        if s > best_score:
            best_score = s
            best_name = f.name
    if best_score >= 2:
        return best_name

    for entry in master_entries:
        entry_blob = _text_blob(entry.name, entry.doc_no.lower())
        if _score_keywords(entry_blob, keywords) < 3:
            continue
        for f in files:
            fn = f.name.lower()
            if _doc_no_matches_filename(entry.doc_no, fn):
                return f.name
    return ""


def _find_checklist_master_entry(
    keywords: list[str],
    master_entries: list[MasterEntry],
) -> str:
    best = ""
    best_score = 0.0
    for entry in master_entries:
        entry_blob = _text_blob(entry.name, entry.doc_no.lower())
        score = _score_keywords(entry_blob, keywords)
        if score > best_score:
            best_score = score
            best = f"{entry.doc_no} - {entry.name}"
    return best if best_score >= 3 else ""


def build_checklist(
    folder: Path,
    image_matches: dict[str, list[ImageMatch]],
    master_path: Path | None,
) -> list[ChecklistItemResult]:
    folder = folder.resolve()
    master_entries: list[MasterEntry] = []
    if master_path and master_path.exists():
        try:
            master_entries = parse_master_list(master_path)
        except Exception:
            pass

    files = list_folder_files(folder)
    if master_path:
        files = [f for f in files if f.resolve() != master_path.resolve()]

    results: list[ChecklistItemResult] = []

    for section, items in CHECKLIST_SECTIONS:
        for label, keywords, photo_slot in items:
            present = False
            matched = ""
            status = "not_found"

            if photo_slot:
                imgs = image_matches.get(photo_slot, [])
                if imgs:
                    present = True
                    matched = ", ".join(m.filename for m in imgs)
                    status = "provided_file"
            elif section == "GENERAL DOCUMENTS" and "Master list" in label:
                if master_path:
                    present = True
                    matched = master_path.name
                    status = "provided_file"
                else:
                    matched = _find_checklist_file(keywords, files, master_entries, folder)
                    present = bool(matched)
                    status = "provided_file" if matched else "not_found"
            else:
                matched = _find_checklist_file(keywords, files, master_entries, folder)
                present = bool(matched)
                status = "provided_file" if matched else "not_found"
                if not matched:
                    master_only = _find_checklist_master_entry(keywords, master_entries)
                    if master_only:
                        present = True
                        matched = master_only
                        status = "master_list_only"

            results.append(
                ChecklistItemResult(
                    section=section,
                    label=label,
                    present=present,
                    matched_file=matched,
                    status=status,
                )
            )

    return results


def format_checklist_text(items: list[ChecklistItemResult]) -> str:
    lines: list[str] = []
    current_section = ""
    for item in items:
        if item.section != current_section:
            if lines:
                lines.append("")
            lines.append(item.section)
            current_section = item.section
        line = item.label
        if item.status == "master_list_only":
            line += " LISTED IN MASTER LIST ONLY"
        elif not item.present:
            line += " NOT THERE"
        lines.append(line)
    return "\n".join(lines) + "\n"


def write_checklist(items: list[ChecklistItemResult], path: Path) -> str:
    text = format_checklist_text(items)
    path.write_text(text, encoding="utf-8")
    return text


def _is_photo_pdf(name: str) -> bool:
    lower = name.lower()
    # Named document PDFs are never factory photos
    if any(x in lower for x in (
        "wi ", " wi-", "sop ", "sop for", "work instruction", "procedure",
        "quality plan", "flow chart", "certificate", "report ", "checklist",
        "master list", "organization",
    )):
        return False
    return any(h in lower for h in PDF_PHOTO_HINTS)


def _pdf_first_page_to_png(path: Path) -> tuple[bytes, str] | None:
    try:
        import fitz  # pymupdf
    except ImportError:
        return None
    try:
        doc = fitz.open(str(path))
        if doc.page_count == 0:
            return None
        pix = doc.load_page(0).get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
        return pix.tobytes("png"), "png"
    except Exception:
        return None


def _load_image_bytes(path: Path) -> tuple[bytes, str] | None:
    ext = path.suffix.lower()
    if ext in IMAGE_EXTENSIONS:
        return path.read_bytes(), ext.lstrip(".")
    if ext == ".pdf" and _is_photo_pdf(path.name):
        return _pdf_first_page_to_png(path)
    return None


def _pdf_all_pages_to_png(path: Path, max_pages: int = 4) -> list[tuple[bytes, str]]:
    """Return PNG bytes for every page of a PDF (up to max_pages)."""
    try:
        import fitz
    except ImportError:
        return []
    pages: list[tuple[bytes, str]] = []
    try:
        doc = fitz.open(str(path))
        for i in range(min(doc.page_count, max_pages)):
            pix = doc.load_page(i).get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
            pages.append((pix.tobytes("png"), "png"))
    except Exception:
        pass
    return pages


def _load_document_preview_bytes(path: Path) -> tuple[bytes, str] | None:
    ext = path.suffix.lower()
    if ext in IMAGE_EXTENSIONS:
        return path.read_bytes(), ext.lstrip(".")
    if ext == ".pdf":
        return _pdf_first_page_to_png(path)
    return None


def _load_document_all_pages(path: Path) -> list[tuple[bytes, str]]:
    """Return preview bytes for every page / image of a document."""
    ext = path.suffix.lower()
    if ext in IMAGE_EXTENSIONS:
        return [(path.read_bytes(), ext.lstrip("."))]
    if ext == ".pdf":
        return _pdf_all_pages_to_png(path) or []
    return []


def _resize_image(data: bytes, ext: str, *, force_jpeg: bool = False) -> tuple[bytes, str, int, int]:
    """Return resized image bytes, normalized ext, width px, height px."""
    img = Image.open(BytesIO(data))
    img = ImageOps.exif_transpose(img)
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
        ext = "jpeg"
    if force_jpeg and img.mode == "L":
        img = img.convert("RGB")
    img.thumbnail((IMAGE_MAX_WIDTH_PX, IMAGE_MAX_WIDTH_PX), Image.Resampling.LANCZOS)
    out = BytesIO()
    save_ext = "jpeg" if force_jpeg or ext in ("jpg", "jpeg") else ext
    if save_ext == "jpeg":
        img.save(out, format="JPEG", quality=82, optimize=True)
    else:
        img.save(out, format=save_ext.upper())
    return out.getvalue(), save_ext, img.width, img.height


def _keyword_in_blob(keyword: str, blob: str) -> bool:
    keyword = keyword.lower().strip()
    if not keyword:
        return False
    if " " in keyword or len(keyword) > 4:
        return keyword in blob
    return re.search(rf"(?<![a-z]){re.escape(keyword)}(?![a-z])", blob) is not None


def _score_image_for_slot(blob: str, slot: PhotoSlot, filename: str = "") -> float:
    score = 0.0
    for kw in slot.keywords:
        if _keyword_in_blob(kw, blob):
            score += max(3, len(kw.split()) * 3)
    # Prefer filename keywords over parent-folder keywords (e.g. Final Product Photos folder)
    # so label / barcode / trademark files don't all collapse into "product shape".
    if filename:
        name_blob = _text_blob(Path(filename).stem, filename.lower())
        for kw in slot.keywords:
            if _keyword_in_blob(kw, name_blob):
                score += max(4, len(kw.split()) * 4)
        # Strong specific product-media cues in the filename itself
        if slot.slot_id == "fp_trademark" and re.search(r"trade[\s_-]*mark|logo|brand", name_blob):
            score += 20
        if slot.slot_id == "fp_barcode" and re.search(r"barcode|qr[\s_-]*code|\bqr\b", name_blob):
            score += 20
        if slot.slot_id == "fp_product_label" and re.search(r"product[\s_-]*label|\blabel\b", name_blob):
            score += 20
        # product_shape should not win on folder name alone when filename is clearly another media type
        if slot.slot_id == "fp_product_shape" and re.search(
            r"barcode|qr[\s_-]*code|trade[\s_-]*mark|product[\s_-]*label|\blabel\b|\blogo\b",
            name_blob,
        ):
            score = 0
    return score


def _extract_photo_pdf_pages_to_slots(
    pdf_path: Path,
) -> dict[str, list[tuple[bytes, str]]]:
    """For a multi-page photo PDF (e.g. TAR Factory Photos.pdf), extract each page
    as an image and map it to the best-matching FA photo slot using in-page text.

    Returns {slot_id: [(image_bytes, ext), ...]}
    """
    try:
        import fitz
    except ImportError:
        return {}

    # Text fragments in page labels → which slot_id(s) they imply
    PAGE_LABEL_SLOTS: list[tuple[re.Pattern, list[str]]] = [
        (re.compile(r"front\s*main\s*building|main\s*building|front\s*building", re.I), ["fp_front"]),
        (re.compile(r"incoming\s*raw\s*material", re.I), ["fp_incoming"]),
        (re.compile(r"production\s*floor", re.I), ["fp_production"]),
        (re.compile(r"packing\s*area", re.I), ["fp_packing"]),
        (re.compile(r"inspection\s*area", re.I), ["fp_inspection"]),
        (re.compile(r"lab\s*area", re.I), ["fp_lab"]),
        (re.compile(r"equipment\s*used\s*for\s*test", re.I), ["fp_lab"]),
        (re.compile(r"calibration\s*labels?\s*on", re.I), ["fp_cal_label"]),
        (re.compile(r"non\s*conform|rejection\s*area", re.I), ["fp_rejection"]),
        (re.compile(r"final\s*product\s*storage", re.I), ["fp_final_storage"]),
        (re.compile(r"product\s*label|image\s*of\s*the\s*product\s*label", re.I), ["fp_product_label"]),
        (re.compile(r"barcode|qr\s*code", re.I), ["fp_barcode"]),
        (re.compile(r"identity\s*and\s*the\s*shape|shape\s*of\s*the\s*product", re.I), ["fp_product_shape"]),
        (re.compile(r"trademark|trade\s*mark", re.I), ["fp_trademark"]),
    ]

    results: dict[str, list[tuple[bytes, str]]] = {}

    try:
        doc = fitz.open(str(pdf_path))
    except Exception:
        return {}

    for page_idx in range(doc.page_count):
        page = doc[page_idx]
        text = page.get_text()

        # Find which slots this page belongs to
        matched_slots: list[str] = []
        for pat, slot_ids in PAGE_LABEL_SLOTS:
            if pat.search(text):
                matched_slots.extend(slot_ids)
        matched_slots = list(dict.fromkeys(matched_slots))  # deduplicate, preserve order

        if not matched_slots:
            continue

        # Render the page to an image
        try:
            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
            img_bytes = pix.tobytes("jpeg")
            ext = "jpeg"
        except Exception:
            continue

        for slot_id in matched_slots:
            results.setdefault(slot_id, []).append((img_bytes, ext))

    return results


def match_images(folder: Path) -> dict[str, list[ImageMatch]]:
    """Match image/PDF-photo files from folder to FA template photo slots.

    For multi-page photo PDFs (e.g. TAR Factory Photos.pdf) the page text is used
    to determine the correct slot per page. Plain image files are scored by filename.
    """
    folder = folder.resolve()
    slot_matches: dict[str, list[ImageMatch]] = {s.slot_id: [] for s in PHOTO_SLOTS}

    # --- Phase 1: page-level extraction from multi-page photo PDFs ---
    page_pdf_slots: set[str] = set()   # track which slot_ids came from page-PDFs
    for path in sorted(folder.rglob("*.pdf"), key=lambda p: p.name.lower()):
        if not path.is_file():
            continue
        if "_cert_previews" in str(path):
            continue
        if not _is_photo_pdf(path.name):
            continue
        per_slot = _extract_photo_pdf_pages_to_slots(path)
        if not per_slot:
            continue
        for slot_id, pages in per_slot.items():
            for img_bytes, ext in pages:
                # Write to a temp file so the rest of the pipeline can read it
                import tempfile
                tmp = tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False)
                tmp.write(img_bytes)
                tmp.close()
                slot = next((s for s in PHOTO_SLOTS if s.slot_id == slot_id), None)
                if slot:
                    slot_matches[slot_id].append(
                        ImageMatch(
                            slot_id=slot_id,
                            slot_label=slot.label,
                            filename=f"{path.stem}_p{path.name}_{slot_id}.{ext}",
                            source_path=tmp.name,
                        )
                    )
                    page_pdf_slots.add(slot_id)

    # --- Phase 2: individual image files (plain images + low-page-count PDFs) ---
    candidates: list[tuple[Path, str]] = []
    for path in folder.rglob("*"):
        if not path.is_file():
            continue
        if "_cert_previews" in str(path):
            continue
        if path.suffix.lower() == ".pdf" and _is_photo_pdf(path.name):
            # Already handled by phase 1 if it produced per-page slot data
            if page_pdf_slots:
                continue
        blob = _path_match_blob(folder, path)
        if any(x in blob for x in IMAGE_EXCLUDE_HINTS):
            continue
        loaded = _load_image_bytes(path)
        if loaded:
            candidates.append((path, blob))

    best_for_file: dict[str, tuple[float, Path, PhotoSlot]] = {}
    for path, blob in candidates:
        photo_pdf = path.suffix.lower() == ".pdf" and _is_photo_pdf(path.name)
        fkey = str(path.resolve())
        for slot in PHOTO_SLOTS:
            if slot.anchor_type != "paragraph":
                continue
            if slot.slot_id == "sec_4.2":
                continue
            s = _score_image_for_slot(blob, slot, path.name)
            if photo_pdf and s > 0:
                s += 2
            if s < 3:
                continue
            prev = best_for_file.get(fkey)
            if prev is None or s > prev[0]:
                best_for_file[fkey] = (s, path, slot)

    by_slot: dict[str, list[tuple[float, Path]]] = {}
    for score, path, slot in best_for_file.values():
        by_slot.setdefault(slot.slot_id, []).append((score, path))

    for slot_id, items in by_slot.items():
        items.sort(key=lambda x: x[1].name.lower())
        slot = next(s for s in PHOTO_SLOTS if s.slot_id == slot_id)
        for _, path in items:
            slot_matches[slot_id].append(
                ImageMatch(
                    slot_id=slot_id,
                    slot_label=slot.label,
                    filename=path.name,
                    source_path=str(path),
                )
            )

    return {k: v for k, v in slot_matches.items() if v}


def _para_text(p: ET.Element) -> str:
    return "".join(t.text or "" for t in p.findall(f".//{W}t"))


def _has_image(elem: ET.Element) -> bool:
    return bool(elem.findall(f".//{A}blip"))


def _factory_photos_bounds(body: ET.Element) -> tuple[int, int] | None:
    """Return (start_idx, end_idx) for the Factory Photos checklist section."""
    children = list(body)
    start = end = None
    for i, child in enumerate(children):
        if not child.tag.endswith("p"):
            continue
        txt = _para_text(child).strip()
        if txt == FACTORY_PHOTOS_START:
            start = i
        elif start is not None and txt == FACTORY_PHOTOS_END:
            end = i
            break
    if start is None or end is None:
        return None
    return start, end


def _clean_factory_photos_images(root: ET.Element) -> None:
    """Remove all image paragraphs in the Factory Photos section (keeps text headings)."""
    body = root.find(f"{W}body")
    if body is None:
        return
    bounds = _factory_photos_bounds(body)
    if not bounds:
        return
    start, end = bounds
    for i in range(end - 1, start, -1):
        child = body[i]
        if not child.tag.endswith("p"):
            continue
        txt = _para_text(child).strip()
        if txt in CERT_HEADING_LABELS:
            continue
        if _has_image(child) or not txt:
            body.remove(child)


def _make_photo_cell(heading: str, image_paras: list[ET.Element]) -> ET.Element:
    """Build a <w:tc> with bold caption and images placed side-by-side inline.

    Multiple images under one heading are placed right next to each other
    in a single paragraph (inline runs), not stacked vertically.
    """
    tc = ET.Element(f"{W}tc")
    tcp = ET.SubElement(tc, f"{W}tcPr")
    tcw = ET.SubElement(tcp, f"{W}tcW")
    tcw.set(f"{W}w", str(PHOTO_CELL_WIDTH))
    tcw.set(f"{W}type", "dxa")
    tcp_mar = ET.SubElement(tcp, f"{W}tcMar")
    for side, pts in (("top", 72), ("bottom", 72), ("left", 108), ("right", 108)):
        m = ET.SubElement(tcp_mar, f"{W}{side}")
        m.set(f"{W}w", str(pts))
        m.set(f"{W}type", "dxa")

    # Heading paragraph
    hp = ET.Element(f"{W}p")
    hp_pr = ET.SubElement(hp, f"{W}pPr")
    jc = ET.SubElement(hp_pr, f"{W}jc")
    jc.set(f"{W}val", "left")
    spc = ET.SubElement(hp_pr, f"{W}spacing")
    spc.set(f"{W}before", "0")
    spc.set(f"{W}after", "40")
    hr = ET.SubElement(hp, f"{W}r")
    hr_pr = ET.SubElement(hr, f"{W}rPr")
    ET.SubElement(hr_pr, f"{W}b")
    sz = ET.SubElement(hr_pr, f"{W}sz")
    sz.set(f"{W}val", "18")
    szcs = ET.SubElement(hr_pr, f"{W}szCs")
    szcs.set(f"{W}val", "18")
    ht = ET.SubElement(hr, f"{W}t")
    ht.text = heading
    tc.append(hp)

    if not image_paras:
        ET.SubElement(tc, f"{W}p")
        return tc

    # All images inline in ONE paragraph — side by side, not stacked
    img_p = ET.Element(f"{W}p")
    img_ppr = ET.SubElement(img_p, f"{W}pPr")
    img_jc = ET.SubElement(img_ppr, f"{W}jc")
    img_jc.set(f"{W}val", "center")

    for ip in image_paras:
        _normalize_photo_size(ip)
        # Extract all <w:r> children from each image paragraph and add to shared paragraph
        for run in ip.findall(f"{W}r"):
            img_p.append(run)

    tc.append(img_p)
    return tc


def _make_empty_photo_cell() -> ET.Element:
    """Empty filler cell to pad the last row of the 2-col photo table."""
    tc = ET.Element(f"{W}tc")
    tcp = ET.SubElement(tc, f"{W}tcPr")
    tcw = ET.SubElement(tcp, f"{W}tcW")
    tcw.set(f"{W}w", str(PHOTO_CELL_WIDTH))
    tcw.set(f"{W}type", "dxa")
    ET.SubElement(tc, f"{W}p")
    return tc


def _make_full_width_photo_cell(heading: str, image_paras: list[ET.Element]) -> ET.Element:
    """Full-width cell (gridSpan=2) with heading + images side by side inline."""
    tc = _make_photo_cell(heading, image_paras)
    tcp = tc.find(f"{W}tcPr")
    if tcp is None:
        tcp = ET.SubElement(tc, f"{W}tcPr")
        tc.insert(0, tcp)
    # Override width to full table width
    tcw = tcp.find(f"{W}tcW")
    if tcw is None:
        tcw = ET.SubElement(tcp, f"{W}tcW")
    tcw.set(f"{W}w", str(PHOTO_TABLE_FULL_WIDTH))
    tcw.set(f"{W}type", "dxa")
    # Span both columns
    span = tcp.find(f"{W}gridSpan")
    if span is None:
        span = ET.SubElement(tcp, f"{W}gridSpan")
    span.set(f"{W}val", "2")
    return tc


def _convert_factory_photos_to_table(root: ET.Element) -> None:
    """Re-layout the Factory Photos section as a borderless 2-column table.

    - One picture per title → titles/side-by-side in 2 columns
    - Multiple pictures per title → full-width row, images stacked one after another under the heading
    """
    body = root.find(f"{W}body")
    if body is None:
        return
    bounds = _factory_photos_bounds(body)
    if not bounds:
        return
    start_idx, end_idx = bounds

    children = list(body)
    items: list[tuple[str, list[ET.Element]]] = []
    current_heading: str | None = None
    current_images: list[ET.Element] = []
    to_remove: list[ET.Element] = []

    for i in range(start_idx + 1, end_idx):
        child = children[i]
        if not child.tag.endswith("p"):
            continue
        txt = _para_text(child).strip()
        has_img = _has_image(child)
        if has_img:
            if current_heading is not None:
                current_images.append(child)
            to_remove.append(child)
        elif txt:
            if current_heading is not None:
                items.append((current_heading, current_images))
            current_heading = txt
            current_images = []
            to_remove.append(child)
        else:
            to_remove.append(child)

    if current_heading is not None:
        items.append((current_heading, current_images))

    if not items:
        return

    for child in to_remove:
        try:
            body.remove(child)
        except ValueError:
            pass

    # Build borderless 2-column table
    tbl = ET.Element(f"{W}tbl")
    tbl_pr = ET.SubElement(tbl, f"{W}tblPr")
    tbl_w = ET.SubElement(tbl_pr, f"{W}tblW")
    tbl_w.set(f"{W}w", str(PHOTO_TABLE_FULL_WIDTH))
    tbl_w.set(f"{W}type", "dxa")
    tbl_layout = ET.SubElement(tbl_pr, f"{W}tblLayout")
    tbl_layout.set(f"{W}type", "fixed")
    tbl_borders = ET.SubElement(tbl_pr, f"{W}tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = ET.SubElement(tbl_borders, f"{W}{side}")
        b.set(f"{W}val", "none")
    tbl_ind = ET.SubElement(tbl_pr, f"{W}tblInd")
    tbl_ind.set(f"{W}w", "0")
    tbl_ind.set(f"{W}type", "dxa")
    tbl_grid = ET.SubElement(tbl, f"{W}tblGrid")
    gc1 = ET.SubElement(tbl_grid, f"{W}gridCol")
    gc1.set(f"{W}w", str(PHOTO_CELL_WIDTH))
    gc2 = ET.SubElement(tbl_grid, f"{W}gridCol")
    gc2.set(f"{W}w", str(PHOTO_CELL_WIDTH))

    idx = 0
    while idx < len(items):
        h1, imgs1 = items[idx]
        # Multiple images under one title → full-width row, stacked one after another
        if len(imgs1) > 1:
            tr = ET.SubElement(tbl, f"{W}tr")
            tr.append(_make_full_width_photo_cell(h1, imgs1))
            idx += 1
            continue

        # Single (or no) image → pair side-by-side with the next single-image title
        tr = ET.SubElement(tbl, f"{W}tr")
        tr.append(_make_photo_cell(h1, imgs1))
        idx += 1
        if idx < len(items) and len(items[idx][1]) <= 1:
            h2, imgs2 = items[idx]
            tr.append(_make_photo_cell(h2, imgs2))
            idx += 1
        else:
            tr.append(_make_empty_photo_cell())

    # Re-locate "Factory Photos:" after removals and insert table right after it
    for i, child in enumerate(body):
        if child.tag.endswith("p") and _para_text(child).strip() == FACTORY_PHOTOS_START:
            body.insert(i + 1, tbl)
            break


def _make_photo_cell_from_rids(
    heading: str,
    images: list[tuple[str, int, int]],
    doc_id_start: int,
    *,
    full_width: bool = False,
) -> tuple[ET.Element, int]:
    """Build a borderless <w:tc>: bold heading + all images inline (side by side).

    images: list of (rel_id, cx_emu, cy_emu). Returns (cell, next_doc_id).
    """
    doc_id = doc_id_start
    tc = ET.Element(f"{W}tc")
    tcp = ET.SubElement(tc, f"{W}tcPr")
    tcw = ET.SubElement(tcp, f"{W}tcW")
    tcw.set(f"{W}w", str(PHOTO_TABLE_FULL_WIDTH if full_width else PHOTO_CELL_WIDTH))
    tcw.set(f"{W}type", "dxa")
    if full_width:
        span = ET.SubElement(tcp, f"{W}gridSpan")
        span.set(f"{W}val", "2")
    tcp_mar = ET.SubElement(tcp, f"{W}tcMar")
    for side, pts in (("top", 72), ("bottom", 72), ("left", 108), ("right", 108)):
        m = ET.SubElement(tcp_mar, f"{W}{side}")
        m.set(f"{W}w", str(pts))
        m.set(f"{W}type", "dxa")

    # Heading
    hp = ET.SubElement(tc, f"{W}p")
    _set_keep_next(hp)
    hp_pr = hp.find(f"{W}pPr")
    if hp_pr is None:
        hp_pr = ET.SubElement(hp, f"{W}pPr")
    ET.SubElement(hp_pr, f"{W}jc").set(f"{W}val", "left")
    spc = ET.SubElement(hp_pr, f"{W}spacing")
    spc.set(f"{W}before", "0")
    spc.set(f"{W}after", "40")
    hr = ET.SubElement(hp, f"{W}r")
    hr_pr = ET.SubElement(hr, f"{W}rPr")
    ET.SubElement(hr_pr, f"{W}b")
    ET.SubElement(hr_pr, f"{W}sz").set(f"{W}val", "18")
    ET.SubElement(hr_pr, f"{W}szCs").set(f"{W}val", "18")
    ET.SubElement(hr, f"{W}t").text = heading

    if not images:
        ET.SubElement(tc, f"{W}p")
        return tc, doc_id

    # All images inline in ONE paragraph — centered within the column so that
    # single-image headings keep the neat 2-column grid.
    img_p = ET.SubElement(tc, f"{W}p")
    img_ppr = ET.SubElement(img_p, f"{W}pPr")
    ET.SubElement(img_ppr, f"{W}jc").set(f"{W}val", "center")
    for rid, cx, cy in images:
        img_p.append(_make_inline_image_run(rid, cx, cy, doc_id))
        doc_id += 1
    return tc, doc_id


def _build_factory_photos_section(
    root: ET.Element,
    section_content: dict[str, list[tuple[str, int, int]]],
) -> None:
    """Deterministically rebuild the whole Factory Photos section as a 2-col table.

    section_content maps a FACTORY_PHOTO_LAYOUT source key to a list of
    (rel_id, cx_emu, cy_emu). Every heading in the layout is always emitted, even
    when it has no images. This replaces the fragile insert-then-reparse pipeline
    and guarantees each heading only shows its own images.
    """
    body = root.find(f"{W}body")
    if body is None:
        return

    children = list(body)
    start_idx = None
    end_idx = len(children)
    for i, child in enumerate(children):
        if not child.tag.endswith("p"):
            continue
        txt = _para_text(child).strip()
        if txt == FACTORY_PHOTOS_START:
            start_idx = i
        elif start_idx is not None and txt.lower().startswith("annexure"):
            end_idx = i
            break
    if start_idx is None:
        return

    # Remove everything between the "Factory Photos:" paragraph and the Annexure
    for child in children[start_idx + 1:end_idx]:
        try:
            body.remove(child)
        except ValueError:
            pass

    # Build the borderless 2-column table
    tbl = ET.Element(f"{W}tbl")
    tbl_pr = ET.SubElement(tbl, f"{W}tblPr")
    tbl_w = ET.SubElement(tbl_pr, f"{W}tblW")
    tbl_w.set(f"{W}w", str(PHOTO_TABLE_FULL_WIDTH))
    tbl_w.set(f"{W}type", "dxa")
    ET.SubElement(tbl_pr, f"{W}tblLayout").set(f"{W}type", "fixed")
    tbl_borders = ET.SubElement(tbl_pr, f"{W}tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        ET.SubElement(tbl_borders, f"{W}{side}").set(f"{W}val", "none")
    tbl_ind = ET.SubElement(tbl_pr, f"{W}tblInd")
    tbl_ind.set(f"{W}w", "0")
    tbl_ind.set(f"{W}type", "dxa")
    tbl_grid = ET.SubElement(tbl, f"{W}tblGrid")
    ET.SubElement(tbl_grid, f"{W}gridCol").set(f"{W}w", str(PHOTO_CELL_WIDTH))
    ET.SubElement(tbl_grid, f"{W}gridCol").set(f"{W}w", str(PHOTO_CELL_WIDTH))

    doc_id = _max_drawing_doc_id(root) + 1
    idx = 0
    layout = [
        item for item in FACTORY_PHOTO_LAYOUT
        if not (item[1] in {"iso9001", "iso14001", "iso45001", "gmp"} and not section_content.get(item[1]))
    ]
    while idx < len(layout):
        heading, key = layout[idx]
        imgs = section_content.get(key, [])
        # Multiple images → full-width row so they sit side by side across the page
        if len(imgs) > 1:
            tr = ET.SubElement(tbl, f"{W}tr")
            _set_cant_split(tr)
            cell, doc_id = _make_photo_cell_from_rids(heading, imgs, doc_id, full_width=True)
            tr.append(cell)
            idx += 1
            continue
        # Single/zero image → pair with the next single/zero-image heading
        tr = ET.SubElement(tbl, f"{W}tr")
        _set_cant_split(tr)
        cell, doc_id = _make_photo_cell_from_rids(heading, imgs, doc_id)
        tr.append(cell)
        idx += 1
        if idx < len(layout) and len(section_content.get(layout[idx][1], [])) <= 1:
            heading2, key2 = layout[idx]
            cell2, doc_id = _make_photo_cell_from_rids(
                heading2, section_content.get(key2, []), doc_id)
            tr.append(cell2)
            idx += 1
        else:
            tr.append(_make_empty_photo_cell())

    for i, child in enumerate(body):
        if child.tag.endswith("p") and _para_text(child).strip() == FACTORY_PHOTOS_START:
            # Drop empty filler paragraphs right before the heading — these leave a
            # blank page between the previous section and the photos.
            j = i - 1
            while j >= 0:
                prev = body[j]
                if (
                    prev.tag.endswith("p")
                    and not _para_text(prev).strip()
                    and prev.find(f".//{W}br") is None
                    and not _has_image(prev)
                ):
                    body.remove(prev)
                    j -= 1
                else:
                    break
            _set_page_break_before(child)
            _set_keep_next(child)
            insert_at = list(body).index(child) + 1
            body.insert(insert_at, tbl)
            break

    for child in body:
        if child.tag.endswith("p") and _para_text(child).strip().lower().startswith("annexure"):
            _set_page_break_before(child)


def _clean_images_in_section_tables(root: ET.Element) -> None:
    """Remove images from checklist table cells (sections 1–8). Signatures are kept."""
    for tbl in root.findall(f".//{W}tbl"):
        for tr in tbl.findall(f"{W}tr"):
            tcs = tr.findall(f"{W}tc")
            if not tcs:
                continue
            row_text = " || ".join(_cell_text(tc) for tc in tcs).lower()
            if re.search(r"signature.*t[üu]v|t[üu]v.*auditor", row_text):
                continue
            if "signature" in row_text and "factory" in row_text:
                continue
            for tc in tcs:
                for p in list(tc.findall(f"{W}p")):
                    if _has_image(p):
                        tc.remove(p)


def _emu_from_px(px: int) -> int:
    return int(px * 9525)


def _normalize_photo_size(
    image_para: ET.Element,
    fixed_cx: int = PHOTO_FIXED_WIDTH_EMU,
    fixed_cy: int = PHOTO_FIXED_HEIGHT_EMU,
) -> None:
    """Rescale inline photo images into a compact uniform box."""
    WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    for inline in image_para.findall(f".//{{{WP_NS}}}inline"):
        extent = inline.find(f"{{{WP_NS}}}extent")
        if extent is not None:
            old_cx = int(extent.get("cx", fixed_cx) or fixed_cx)
            old_cy = int(extent.get("cy", fixed_cx) or fixed_cx)
            ratio = old_cy / max(old_cx, 1)
            cx = fixed_cx
            cy = int(cx * ratio)
            if cy > fixed_cy:
                cy = fixed_cy
                cx = int(cy / max(ratio, 0.01))
            extent.set("cx", str(cx))
            extent.set("cy", str(cy))


def _make_image_paragraph(r_id: str, cx: int, cy: int, doc_id: int) -> ET.Element:
    p = ET.Element(f"{W}p")
    p_pr = ET.SubElement(p, f"{W}pPr")
    jc = ET.SubElement(p_pr, f"{W}jc")
    jc.set(f"{W}val", "center")

    r = ET.SubElement(p, f"{W}r")
    drawing = ET.SubElement(r, f"{W}drawing")

    inline = ET.SubElement(drawing, f"{WP}inline")
    inline.set("distT", "0")
    inline.set("distB", "0")
    inline.set("distL", "0")
    inline.set("distR", "0")

    extent = ET.SubElement(inline, f"{WP}extent")
    extent.set("cx", str(cx))
    extent.set("cy", str(cy))

    doc_pr = ET.SubElement(inline, f"{WP}docPr")
    doc_pr.set("id", str(doc_id))
    doc_pr.set("name", f"FA Photo {doc_id}")

    c_nv = ET.SubElement(inline, f"{WP}cNvGraphicFramePr")
    locks = ET.SubElement(c_nv, f"{A}graphicFrameLocks")
    locks.set("noChangeAspect", "1")

    graphic = ET.SubElement(inline, f"{A}graphic")
    gdata = ET.SubElement(graphic, f"{A}graphicData")
    gdata.set("uri", "http://schemas.openxmlformats.org/drawingml/2006/picture")

    pic = ET.SubElement(gdata, f"{PIC}pic")
    nv = ET.SubElement(pic, f"{PIC}nvPicPr")
    c_nv_pr = ET.SubElement(nv, f"{PIC}cNvPr")
    c_nv_pr.set("id", "0")
    c_nv_pr.set("name", "")
    c_nv_pic = ET.SubElement(nv, f"{PIC}cNvPicPr")
    pic_locks = ET.SubElement(c_nv_pic, f"{A}picLocks")
    pic_locks.set("noChangeAspect", "1")

    blip_fill = ET.SubElement(pic, f"{PIC}blipFill")
    blip = ET.SubElement(blip_fill, f"{A}blip")
    blip.set(f"{R}embed", r_id)
    stretch = ET.SubElement(blip_fill, f"{A}stretch")
    ET.SubElement(stretch, f"{A}fillRect")

    sp_pr = ET.SubElement(pic, f"{PIC}spPr")
    xfrm = ET.SubElement(sp_pr, f"{A}xfrm")
    ET.SubElement(xfrm, f"{A}off", x="0", y="0")
    ext = ET.SubElement(xfrm, f"{A}ext")
    ext.set("cx", str(cx))
    ext.set("cy", str(cy))
    prst = ET.SubElement(sp_pr, f"{A}prstGeom")
    prst.set("prst", "rect")
    ET.SubElement(prst, f"{A}avLst")

    return p


def _make_inline_image_run(r_id: str, cx: int, cy: int, doc_id: int) -> ET.Element:
    """Single image run — used to build multi-image cert paragraphs."""
    r = ET.Element(f"{W}r")
    drawing = ET.SubElement(r, f"{W}drawing")

    inline = ET.SubElement(drawing, f"{WP}inline")
    inline.set("distT", "0")
    inline.set("distB", "0")
    inline.set("distL", "0")
    inline.set("distR", "0")

    extent = ET.SubElement(inline, f"{WP}extent")
    extent.set("cx", str(cx))
    extent.set("cy", str(cy))

    doc_pr = ET.SubElement(inline, f"{WP}docPr")
    doc_pr.set("id", str(doc_id))
    doc_pr.set("name", f"FA Cert {doc_id}")

    c_nv = ET.SubElement(inline, f"{WP}cNvGraphicFramePr")
    locks = ET.SubElement(c_nv, f"{A}graphicFrameLocks")
    locks.set("noChangeAspect", "1")

    graphic = ET.SubElement(inline, f"{A}graphic")
    gdata = ET.SubElement(graphic, f"{A}graphicData")
    gdata.set("uri", "http://schemas.openxmlformats.org/drawingml/2006/picture")

    pic = ET.SubElement(gdata, f"{PIC}pic")
    nv = ET.SubElement(pic, f"{PIC}nvPicPr")
    c_nv_pr = ET.SubElement(nv, f"{PIC}cNvPr")
    c_nv_pr.set("id", "0")
    c_nv_pr.set("name", "")
    c_nv_pic = ET.SubElement(nv, f"{PIC}cNvPicPr")
    pic_locks = ET.SubElement(c_nv_pic, f"{A}picLocks")
    pic_locks.set("noChangeAspect", "1")

    blip_fill = ET.SubElement(pic, f"{PIC}blipFill")
    blip = ET.SubElement(blip_fill, f"{A}blip")
    blip.set(f"{R}embed", r_id)
    stretch = ET.SubElement(blip_fill, f"{A}stretch")
    ET.SubElement(stretch, f"{A}fillRect")

    sp_pr = ET.SubElement(pic, f"{PIC}spPr")
    xfrm = ET.SubElement(sp_pr, f"{A}xfrm")
    ET.SubElement(xfrm, f"{A}off", x="0", y="0")
    ext_el = ET.SubElement(xfrm, f"{A}ext")
    ext_el.set("cx", str(cx))
    ext_el.set("cy", str(cy))
    prst = ET.SubElement(sp_pr, f"{A}prstGeom")
    prst.set("prst", "rect")
    ET.SubElement(prst, f"{A}avLst")
    return r


def _make_cert_image_paragraph(
    rids_wh: list[tuple[str, int, int]],
    doc_id_start: int,
) -> tuple[ET.Element, int]:
    """One paragraph with cert page images side-by-side after an ISO heading."""
    p = ET.Element(f"{W}p")
    p_pr = ET.SubElement(p, f"{W}pPr")
    jc = ET.SubElement(p_pr, f"{W}jc")
    jc.set(f"{W}val", "left")
    doc_id = doc_id_start
    for rid, w, h in rids_wh:
        cx = min(_emu_from_px(w), SUPPORT_DOC_MAX_WIDTH_EMU)
        cy = min(int(cx * h / max(w, 1)), SUPPORT_DOC_MAX_HEIGHT_EMU)
        p.append(_make_inline_image_run(rid, cx, cy, doc_id))
        doc_id += 1
    return p, doc_id


def _max_drawing_doc_id(root: ET.Element) -> int:
    ids = [int(m) for m in re.findall(r'\bdocPr id="(\d+)"', ET.tostring(root, encoding="unicode"))]
    return max(ids) if ids else 100


def _next_rel_ids(rels_xml: bytes, count: int) -> tuple[list[str], int]:
    root = ET.fromstring(rels_xml)
    max_id = 0
    for rel in root:
        rid = rel.get("Id", "")
        m = re.match(r"rId(\d+)", rid)
        if m:
            max_id = max(max_id, int(m.group(1)))
    ids = [f"rId{max_id + i + 1}" for i in range(count)]
    return ids, max_id + count


def _next_fa_media_index(other_files: dict[str, bytes], prefix: str) -> int:
    max_idx = 0
    pattern = re.compile(rf"word/media/{re.escape(prefix)}(\d{{3}})\.")
    for name in other_files:
        m = pattern.search(name)
        if m:
            max_idx = max(max_idx, int(m.group(1)))
    return max_idx + 1


def _add_image_relationships(rels_xml: bytes, items: list[tuple[str, str]]) -> bytes:
    """items: [(rId, media_target)] e.g. ('rId20', 'media/fa_img_001.jpeg')"""
    root = ET.fromstring(rels_xml)
    for rid, target in items:
        rel = ET.SubElement(root, "Relationship")
        rel.set("Id", rid)
        rel.set(
            "Type",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
        )
        rel.set("Target", target)
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _insert_images_in_document(
    root: ET.Element,
    image_payloads: list[tuple[PhotoSlot, bytes, str, int, int, str]],
    *,
    skip_if_present: bool = False,
    section_bounds: tuple[int, int] | None = None,
) -> None:
    """Insert centered, resized images after matching anchors in the template."""
    slot_to_payload: dict[str, list[tuple[bytes, str, int, int, str]]] = {}
    for slot, data, ext, w, h, rid in image_payloads:
        slot_to_payload.setdefault(slot.slot_id, []).append((data, ext, w, h, rid))

    body = root.find(f"{W}body")
    if body is None:
        return

    insert_after: list[tuple[int, ET.Element]] = []
    doc_id = 100

    for idx, child in enumerate(list(body)):
        if not child.tag.endswith("p"):
            continue
        if section_bounds and not (section_bounds[0] < idx < section_bounds[1]):
            continue
        text = _para_text(child)
        for slot in PHOTO_SLOTS:
            if slot.anchor_type != "paragraph":
                continue
            payloads = slot_to_payload.get(slot.slot_id)
            if not payloads:
                continue
            if slot.anchor.lower() not in text.lower():
                continue
            pos = idx + 1
            if skip_if_present and pos < len(body) and _has_image(body[pos]):
                break
            for _, _, w, h, rid in payloads:
                cx = min(_emu_from_px(w), IMAGE_MAX_WIDTH_EMU)
                cy = int(cx * h / max(w, 1))
                insert_after.append((pos, _make_image_paragraph(rid, cx, cy, doc_id)))
                doc_id += 1
                pos += 1
            break

    for pos, elem in sorted(insert_after, key=lambda x: -x[0]):
        body.insert(pos, elem)

    # Photos only go in the Factory Photos checklist (after 9.0), not in section table rows.


def _extract_calibration_cert_number(folder: Path | None, matches: list[DocumentMatch]) -> str:
    if folder:
        for path in sorted(folder.rglob("*"), key=lambda p: p.name.lower()):
            if path.suffix.lower() != ".pdf":
                continue
            if "calibration register" not in path.name.lower():
                continue
            text = _pdf_text(path, max_pages=10)
            nums = re.findall(r"\b([A-Z]{1,4}\d[\d./-]{4,}|\d{10,})\b", text)
            for num in nums:
                if re.match(r"^202[0-9]$", num):
                    continue
                return num.strip()

    for m in matches:
        blob = _text_blob(m.heading, m.filename, m.doc_no)
        for pat in (
            r"\b([A-Z]{1,4}\d[\d./-]{4,})\b",
            r"\b(\d{10,})\b",
        ):
            hit = re.search(pat, blob, re.I)
            if hit:
                return hit.group(1).strip()
        if m.doc_no:
            return m.doc_no
    return ""


# Sections about non-conformity — their document number may be reused elsewhere.
NON_CONFORMITY_SIDS = frozenset({"2.2", "3.5", "4.5"})


def _resolve_match_doc_no(m: DocumentMatch, folder: Path | None) -> str:
    """Doc number from master list / filename prefix / document content."""
    doc_no = m.doc_no or extract_doc_no_from_filename(m.filename or m.heading)
    if not doc_no and m.filename and folder:
        candidate_path = folder / m.filename
        if not candidate_path.exists():
            candidates = list(folder.rglob(m.filename))
            candidate_path = candidates[0] if candidates else candidate_path
        if candidate_path.exists():
            doc_no = _extract_doc_no_from_content(candidate_path)
    return doc_no


def _format_section_rhs(
    sid: str,
    matches: list[DocumentMatch],
    folder: Path | None = None,
    used_doc_nos: set[str] | None = None,
) -> str:
    """One or two compact document refs for the RHS cell under Yes/No.

    Document numbers are not repeated across sections (tracked in used_doc_nos),
    except for non-conformity sections which are allowed to reuse the same doc.
    """
    allow_repeat = sid in NON_CONFORMITY_SIDS
    if used_doc_nos is None:
        used_doc_nos = set()

    if sid == "4.1":
        cert_no = _extract_calibration_cert_number(folder, matches)
        if not cert_no and matches:
            cert_no = Path(matches[0].filename).stem if matches[0].filename else matches[0].heading
        if cert_no:
            if not allow_repeat:
                used_doc_nos.add(cert_no)
            return f"{cert_no}\ncalibration certificate reviewed and satisfactory"
        return "NOT THERE"

    if not matches:
        return "NOT THERE"
    lines: list[str] = []
    for m in matches[:2]:
        doc_no = _resolve_match_doc_no(m, folder)
        clean_name = _clean_doc_heading(m.heading, m.filename)
        # Skip documents already listed under an earlier section (no repeats)
        dedup_key = (doc_no or clean_name or "").strip().lower()
        if dedup_key and not allow_repeat and dedup_key in used_doc_nos:
            continue
        if dedup_key and not allow_repeat:
            used_doc_nos.add(dedup_key)
        if doc_no and clean_name:
            lines.append(f"{doc_no}: {clean_name}")
        elif doc_no:
            lines.append(doc_no)
        elif clean_name:
            lines.append(f"{clean_name}\nReviewed and Satisfactory")
    if not lines:
        return "NOT THERE"
    return "\n".join(lines)


def _find_doc_number_for(folder: Path, patterns: list[str]) -> str:
    """Find the document number of the first file matching any pattern."""
    exts = {".pdf", ".docx", ".doc"} | IMAGE_EXTENSIONS
    for p in sorted(folder.rglob("*"), key=lambda x: (len(x.parts), x.name.lower())):
        if not p.is_file() or p.suffix.lower() not in exts:
            continue
        name = p.name.lower()
        if any(re.search(pat, name) for pat in patterns):
            no = extract_doc_no_from_filename(p.name) or _extract_doc_no_from_content(p)
            if no:
                return no
    return ""


def _fill_supporting_doc_numbers(root: ET.Element, folder: Path | None) -> None:
    """Write the bare document numbers for Company quality manual and
    Control of Non-conforming Product in section 1.4 (no 'DOC No-' prefix)."""
    if not folder:
        return
    targets = {
        "company quality manual": _find_doc_number_for(
            folder, [r"quality manual", r"ims.*manual", r"master list", r"quality\s*manual"]),
        "control of non-conforming": _find_doc_number_for(
            folder, [r"non.?conform", r"control of non"]),
    }
    for tbl in root.findall(f".//{W}tbl"):
        for tr in tbl.findall(f"{W}tr"):
            tcs = tr.findall(f"{W}tc")
            if len(tcs) < 2:
                continue
            label = _cell_text(tcs[0]).strip().lower()
            for key, number in targets.items():
                if not number:
                    continue
                if label.startswith(key):
                    rhs = tcs[1]
                    # Keep the Yes/No checkbox line, replace any trailing doc-no text
                    checkbox_line = _cell_text(rhs).splitlines()[0] if _cell_text(rhs) else ""
                    checkbox_line = re.split(r"\s{2,}(?:doc\s*no|document\s*no)", checkbox_line, flags=re.I)[0].rstrip()
                    _set_cell_value(rhs, f"{checkbox_line}    {number}")
                    break


def _cell_text(tc: ET.Element) -> str:
    parts = []
    for p in tc.findall(f"{W}p"):
        parts.append("".join(t.text or "" for t in p.findall(f".//{W}t")))
    return "\n".join(parts).strip()


def _append_paragraph(
    tc: ET.Element,
    text: str,
    *,
    small: bool = False,
    center: bool = False,
) -> None:
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        p = ET.SubElement(tc, f"{W}p")
        p_pr = ET.SubElement(p, f"{W}pPr")
        if center:
            jc = ET.SubElement(p_pr, f"{W}jc")
            jc.set(f"{W}val", "center")
        spacing = ET.SubElement(p_pr, f"{W}spacing")
        spacing.set(f"{W}before", "0")
        spacing.set(f"{W}after", "20")
        spacing.set(f"{W}line", "240")
        spacing.set(f"{W}lineRule", "auto")
        r = ET.SubElement(p, f"{W}r")
        if small:
            r_pr = ET.SubElement(r, f"{W}rPr")
            sz = ET.SubElement(r_pr, f"{W}sz")
            sz.set(f"{W}val", "16")  # 8pt
            sz_cs = ET.SubElement(r_pr, f"{W}szCs")
            sz_cs.set(f"{W}val", "16")
        t = ET.SubElement(r, f"{W}t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = line


YES_NO_RE = re.compile(r"Yes\s*.?.?.?\s*No", re.I)


def _para_plain(p: ET.Element) -> str:
    return "".join(t.text or "" for t in p.findall(f".//{W}t"))


def _update_rhs_cell(tc: ET.Element, doc_text: str) -> None:
    """Keep Yes/No line, replace document lines underneath with compact text."""
    paragraphs = list(tc.findall(f"{W}p"))
    header_end = 0
    for i, p in enumerate(paragraphs):
        if YES_NO_RE.search(_para_plain(p)):
            header_end = i + 1
            break

    for p in paragraphs[header_end:]:
        tc.remove(p)

    if doc_text:
        _append_paragraph(tc, doc_text, center=True)


def _set_cell_value(tc: ET.Element, value: str) -> None:
    for p in list(tc.findall(f"{W}p")):
        tc.remove(p)
    if value:
        _append_paragraph(tc, value, small=True)


def _set_page_break_before(p: ET.Element) -> None:
    p_pr = p.find(f"{W}pPr")
    if p_pr is None:
        p_pr = ET.Element(f"{W}pPr")
        p.insert(0, p_pr)
    if p_pr.find(f"{W}pageBreakBefore") is None:
        ET.SubElement(p_pr, f"{W}pageBreakBefore")


def _set_keep_next(p: ET.Element) -> None:
    p_pr = p.find(f"{W}pPr")
    if p_pr is None:
        p_pr = ET.Element(f"{W}pPr")
        p.insert(0, p_pr)
    if p_pr.find(f"{W}keepNext") is None:
        ET.SubElement(p_pr, f"{W}keepNext")


def _set_cant_split(tr: ET.Element) -> None:
    tr_pr = tr.find(f"{W}trPr")
    if tr_pr is None:
        tr_pr = ET.Element(f"{W}trPr")
        tr.insert(0, tr_pr)
    if tr_pr.find(f"{W}cantSplit") is None:
        ET.SubElement(tr_pr, f"{W}cantSplit")


def _yes_no_text(label: str, present: bool, include_na: bool = True) -> str:
    if include_na:
        return f"{label}\nYes  {CHECKED if present else UNCHECKED}  No  {UNCHECKED if present else CHECKED} N/A {UNCHECKED}"
    return f"{label}\nYes  {CHECKED if present else UNCHECKED}  No  {UNCHECKED if present else CHECKED}"


def _normalize_label(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip().lower()


def fill_factory_info(root: ET.Element, info: dict[str, str]) -> int:
    """Fill template sections 1.1–1.3 from Factory Info PDF."""
    if not info:
        return 0
    filled = 0
    label_to_key: dict[str, str] = {}
    for key, labels in TEMPLATE_FACTORY_LABELS.items():
        for label in labels:
            label_to_key[_normalize_label(label)] = key

    for tbl in root.findall(f".//{W}tbl"):
        for tr in tbl.findall(f"{W}tr"):
            tcs = tr.findall(f"{W}tc")
            if len(tcs) < 2:
                continue
            label = _normalize_label(_cell_text(tcs[0]))
            key = label_to_key.get(label)
            if key and info.get(key):
                _set_cell_value(tcs[1], info[key])
                filled += 1
            # Also check right-hand side of 4-column rows (e.g. Factory Area | employees)
            if len(tcs) >= 4:
                label2 = _normalize_label(_cell_text(tcs[2]))
                key2 = label_to_key.get(label2)
                if key2 and info.get(key2):
                    _set_cell_value(tcs[3], info[key2])
                    filled += 1

        for tr in tbl.findall(f"{W}tr"):
            tcs = tr.findall(f"{W}tc")
            if len(tcs) == 2 and _normalize_label(_cell_text(tcs[0])) == "factory name:":
                name = info.get("registered_name", "")
                if name:
                    _set_cell_value(tcs[1], name)
                    filled += 1
                break

    return filled


def fill_audit_meta(root: ET.Element, meta: dict[str, Any]) -> int:
    filled = 0
    factory_name = str(meta.get("factory_name", "")).strip()
    regulation = str(meta.get("technical_regulation", "")).strip()
    audit_date = str(meta.get("audit_date", "")).strip()
    auditor_names = meta.get("auditor_names", [])
    if isinstance(auditor_names, str):
        auditor_names = [auditor_names]
    auditor_text = ", ".join(str(n).strip() for n in auditor_names if str(n).strip())
    factory_rep = str(meta.get("factory_representative", "")).strip()
    report_number = str(meta.get("report_number", "")).strip()
    if not report_number and meta.get("country"):
        report_number = generate_report_number(
            str(meta.get("country", "")),
            auditor_names if isinstance(auditor_names, list) else [],
            audit_date,
        )
    support_docs = meta.get("support_docs", {}) if isinstance(meta.get("support_docs", {}), dict) else {}

    for tbl in root.findall(f".//{W}tbl"):
        for tr in tbl.findall(f"{W}tr"):
            tcs = tr.findall(f"{W}tc")
            if not tcs:
                continue
            row_text = " || ".join(_cell_text(tc) for tc in tcs)
            first = _normalize_label(_cell_text(tcs[0]))

            if report_number and len(tcs) >= 2 and first == "report no.:":
                _set_cell_value(tcs[1], report_number)
                filled += 1

            if len(tcs) >= 2 and first == "factory name:":
                _set_cell_value(tcs[1], factory_name)
                filled += 1

            if "technical regulation" in row_text.lower():
                target = tcs[1] if len(tcs) >= 2 else tcs[0]
                reg_text = regulation
                if re.search(r"\bSASO\s+Saleem\s+Program\b", reg_text, re.I):
                    reg_text = ""
                if not reg_text.lower().startswith("technical regulation"):
                    reg_text = f"Technical Regulation for {reg_text}".rstrip()
                _set_cell_value(target, reg_text)
                filled += 1

            if audit_date and "the audit was carried out" in row_text.lower():
                _set_cell_value(tcs[0], f"The audit was carried out (between): {audit_date}")
                filled += 1

            if audit_date and len(tcs) == 2 and all(
                _normalize_label(_cell_text(tc)).startswith("date") for tc in tcs
            ):
                _set_cell_value(tcs[0], f"Date: {audit_date}")
                _set_cell_value(tcs[1], f"Date: {audit_date}")
                filled += 2

            if audit_date and len(tcs) >= 4 and first == "date" and _normalize_label(_cell_text(tcs[2])) == "date":
                _set_cell_value(tcs[1], audit_date)
                _set_cell_value(tcs[3], audit_date)
                filled += 2

            if "copy of certificates" in row_text.lower() and len(tcs) >= 5:
                _set_cell_value(tcs[1], _yes_no_text("ISO 9001", bool(support_docs.get("iso9001"))))
                _set_cell_value(tcs[2], _yes_no_text("ISO 14001", bool(support_docs.get("iso14001"))))
                _set_cell_value(tcs[3], _yes_no_text("ISO 45001", bool(support_docs.get("iso45001"))))
                _set_cell_value(tcs[4], _yes_no_text("GMP", bool(support_docs.get("gmp"))))
                filled += 4

            if "if the manufacturer operates a quality management system" in row_text.lower():
                iso_details = support_docs.get("iso9001_details", {}) if isinstance(support_docs, dict) else {}
                if support_docs.get("iso9001"):
                    qms_text = (
                        f"If the manufacturer operates a Quality Management System, which kind of?\n"
                        f"{UNCHECKED} Quality Management System NOT certified\n"
                        f"{CHECKED} Quality Management System certified by an accredited Body\n"
                        f"{UNCHECKED} Quality Management System certified by a non-accredited Body\n"
                        f"{CHECKED} Copy of the certificate provided as appendix to this report\n"
                        f"If the QMS is certified or assessed by an accredited or non-accredited Body, "
                        f"provide details and a copy of the certificate.\n"
                        f"Details of QMS standard: {iso_details.get('standard') or 'ISO 9001'}\n"
                        f"Name of certification body: {iso_details.get('body') or 'NA'}\n"
                        f"Certificate no: {iso_details.get('cert_no') or 'NA'}\n"
                        f"Certificate issued date/Certificate expiry date: {iso_details.get('dates') or 'NA'}"
                    )
                else:
                    qms_text = (
                        f"If the manufacturer operates a Quality Management System, which kind of?\n"
                        f"{CHECKED} Quality Management System NOT certified\n"
                        f"{UNCHECKED} Quality Management System certified by an accredited Body\n"
                        f"{UNCHECKED} Quality Management System certified by a non-accredited Body\n"
                        f"{UNCHECKED} Copy of the certificate provided as appendix to this report\n"
                        f"If the QMS is certified or assessed by an accredited or non-accredited Body, "
                        f"provide details and a copy of the certificate.\n"
                        f"Details of QMS standard: NA\n"
                        f"Name of certification body: NA\n"
                        f"Certificate no: NA\n"
                        f"Certificate issued date/Certificate expiry date: NA"
                    )
                _set_cell_value(tcs[0], qms_text)
                filled += 1

            if len(tcs) == 2 and "auditor" in _cell_text(tcs[0]).lower() and "printed letters" in _cell_text(tcs[0]).lower():
                _set_cell_value(tcs[0], f"Auditor’s name (printed letters): {auditor_text}" if auditor_text else "Auditor’s name (printed letters):")
                _set_cell_value(tcs[1], f"Factory representative (printed letters): {factory_rep}" if factory_rep else "Factory representative (printed letters):")
                filled += 2

    # Fill plain-paragraph "Date:" and "Location:" on the cover page
    body = root.find(f"{W}body")
    if body is not None:
        for child in body:
            if not child.tag.endswith("p"):
                continue
            txt = _para_text(child).strip().lower()
            if audit_date and txt in ("date:", "date"):
                _set_para_text(child, f"Date: {audit_date}")
                filled += 1
            elif txt in ("location:", "location"):
                # Location shows the country only
                loc = meta.get("country") or ""
                if loc:
                    _set_para_text(child, f"Location: {loc}")
                    filled += 1

    return filled


def _set_para_text(p: ET.Element, text: str) -> None:
    """Replace all runs in a paragraph with a single run containing text."""
    for r in list(p.findall(f"{W}r")):
        p.remove(r)
    r = ET.SubElement(p, f"{W}r")
    t = ET.SubElement(r, f"{W}t")
    t.text = text
    if text.startswith(" ") or text.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def _clear_signature_cell(tc: ET.Element, label: str = "Signature:") -> None:
    for p in list(tc.findall(f"{W}p")):
        tc.remove(p)
    _append_paragraph(tc, label)


def _append_signature_images(
    tc: ET.Element,
    payloads: list[tuple[str, int, int, str]],
    doc_id_start: int,
) -> int:
    doc_id = doc_id_start
    for _, w, h, rid in payloads:
        cx = min(_emu_from_px(w), SIGNATURE_MAX_WIDTH_EMU)
        cy = int(cx * h / max(w, 1))
        tc.append(_make_image_paragraph(rid, cx, cy, doc_id))
        doc_id += 1
    return doc_id


def _insert_signatures_in_document(
    root: ET.Element,
    signature_payloads: list[tuple[str, int, int, str]],
    *,
    skip_if_present: bool = False,
) -> None:
    auditor_payloads = [p for p in signature_payloads if p[0] == "auditor"]
    cover_auditor_payloads = [p for p in signature_payloads if p[0] == "auditor_cover"]
    factory_payloads = [p for p in signature_payloads if p[0] == "factory"]
    doc_id = 600

    for tbl in root.findall(f".//{W}tbl"):
        for tr in tbl.findall(f"{W}tr"):
            tcs = tr.findall(f"{W}tc")
            if not tcs:
                continue
            row_text = " || ".join(_cell_text(tc) for tc in tcs).lower()

            if "signature, tüv rheinland auditor" in row_text or "signature, tuv rheinland auditor" in row_text:
                if len(tcs) >= 2:
                    if not (skip_if_present and _has_image(tcs[1])):
                        _clear_signature_cell(tcs[1], "")
                        doc_id = _append_signature_images(tcs[1], cover_auditor_payloads or auditor_payloads[:1], doc_id)
                if len(tcs) >= 3:
                    if not (skip_if_present and _has_image(tcs[2])):
                        _clear_signature_cell(tcs[2], "Signature, Factory representative")
                        doc_id = _append_signature_images(tcs[2], factory_payloads, doc_id)

            if len(tcs) == 2 and all(_cell_text(tc).strip().lower().startswith("signature") for tc in tcs):
                if not (skip_if_present and _has_image(tcs[0])):
                    _clear_signature_cell(tcs[0], "Signature:")
                    doc_id = _append_signature_images(tcs[0], auditor_payloads, doc_id)
                if not (skip_if_present and _has_image(tcs[1])):
                    _clear_signature_cell(tcs[1], "Signature:")
                    doc_id = _append_signature_images(tcs[1], factory_payloads, doc_id)

            if len(tcs) >= 4 and _normalize_label(_cell_text(tcs[0])) == "signature" and _normalize_label(_cell_text(tcs[2])) == "signature":
                if not (skip_if_present and _has_image(tcs[1])):
                    _clear_signature_cell(tcs[1], "")
                    doc_id = _append_signature_images(tcs[1], factory_payloads, doc_id)
                if not (skip_if_present and _has_image(tcs[3])):
                    _clear_signature_cell(tcs[3], "")
                    doc_id = _append_signature_images(tcs[3], auditor_payloads, doc_id)


def _insert_support_documents_in_document(
    root: ET.Element,
    support_payloads: list[tuple[str, int, int, str]],
    *,
    skip_if_present: bool = False,
    section_bounds: tuple[int, int] | None = None,
) -> None:
    """Insert cert images in a new paragraph directly after each ISO/GMP heading."""
    if not support_payloads:
        return

    CERT_HEADINGS: dict[str, str] = {
        "iso9001": "ISO 9001",
        "iso14001": "ISO 14001",
        "iso45001": "ISO 45001",
        "gmp": "GMP",
    }
    heading_texts = set(CERT_HEADINGS.values())

    payload_by_key: dict[str, list[tuple[str, int, int]]] = {}
    for key, w, h, rid in support_payloads:
        payload_by_key.setdefault(key, []).append((rid, w, h))

    body = root.find(f"{W}body")
    if body is None:
        return

    doc_id = _max_drawing_doc_id(root) + 1
    insertions: list[tuple[int, ET.Element]] = []

    children = list(body)
    for idx, child in enumerate(children):
        if not child.tag.endswith("p"):
            continue
        if section_bounds and not (section_bounds[0] < idx < section_bounds[1]):
            continue
        txt = _para_text(child).strip()
        for cert_key, heading in CERT_HEADINGS.items():
            if txt != heading:
                continue
            rids_wh = payload_by_key.get(cert_key)
            if not rids_wh:
                break

            # Remove stale cert-image paragraphs already sitting after this heading
            remove_at = idx + 1
            while remove_at < len(body):
                nxt = body[remove_at]
                if not nxt.tag.endswith("p"):
                    break
                ntxt = _para_text(nxt).strip()
                if ntxt in heading_texts:
                    break
                if _has_image(nxt) or not ntxt:
                    body.remove(nxt)
                    continue
                break

            if skip_if_present:
                nxt = body[idx + 1] if idx + 1 < len(body) else None
                if nxt is not None and nxt.tag.endswith("p") and _has_image(nxt):
                    break

            img_p, doc_id = _make_cert_image_paragraph(rids_wh, doc_id)
            insertions.append((idx + 1, img_p))
            break

    for pos, elem in sorted(insertions, key=lambda x: -x[0]):
        body.insert(pos, elem)


def _insert_cert_images_docx(
    docx_path: Path,
    support_items: list[tuple[str, Path]],
) -> None:
    """Insert ISO/GMP certificate images after heading paragraphs using python-docx."""
    if not support_items:
        return
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.shared import Inches
        from docx.text.paragraph import Paragraph
    except ImportError:
        return

    CERT_HEADINGS: dict[str, str] = {
        "iso9001": "ISO 9001",
        "iso14001": "ISO 14001",
        "iso45001": "ISO 45001",
        "gmp": "GMP",
    }
    heading_texts = set(CERT_HEADINGS.values())
    blip_tag = f"{{{A_NS}}}blip"

    preview_dir = docx_path.parent / "_cert_previews"
    preview_dir.mkdir(exist_ok=True)
    images_by_heading: dict[str, list[Path]] = {}

    for key, doc_path in support_items:
        heading = CERT_HEADINGS.get(key)
        if not heading:
            continue
        page_paths: list[Path] = []
        for page_i, (page_data, page_ext) in enumerate(_load_document_all_pages(doc_path)):
            data, ext, _, _ = _resize_image(page_data, page_ext, force_jpeg=True)
            out_path = preview_dir / f"{key}_p{page_i + 1}.{ext}"
            out_path.write_bytes(data)
            page_paths.append(out_path)
        if page_paths:
            images_by_heading[heading] = page_paths

    if not images_by_heading:
        return

    def _para_has_image(para: Paragraph) -> bool:
        return bool(para._element.findall(f".//{blip_tag}"))

    def _insert_para_after(para: Paragraph) -> Paragraph:
        new_p = OxmlElement("w:p")
        para._element.addnext(new_p)
        return Paragraph(new_p, para._parent)

    def _remove_para(para: Paragraph) -> None:
        parent = para._element.getparent()
        if parent is not None:
            parent.remove(para._element)

    doc = Document(str(docx_path))

    # Only place certs under the Factory Photos checklist headings (not annexures/tables)
    start_idx = end_idx = None
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if t == "Factory Photos:":
            start_idx = i
        elif start_idx is not None and t == "Factory Layout":
            end_idx = i
            break

    if start_idx is None or end_idx is None:
        doc.save(str(docx_path))
        return

    for i, para in enumerate(doc.paragraphs):
        if not (start_idx < i < end_idx):
            continue
        heading = para.text.strip()
        if heading not in images_by_heading:
            continue

        # Drop any image-only paragraph already sitting after this heading
        next_el = para._element.getnext()
        while next_el is not None and next_el.tag.endswith("p"):
            next_para = Paragraph(next_el, para._parent)
            next_text = next_para.text.strip()
            if next_text in heading_texts:
                break
            if _para_has_image(next_para) or not next_text:
                to_remove = next_para
                next_el = next_el.getnext()
                _remove_para(to_remove)
                continue
            break

        img_para = _insert_para_after(para)
        for img_path in images_by_heading[heading]:
            run = img_para.add_run()
            run.add_picture(str(img_path), width=Inches(1.85))

    doc.save(str(docx_path))


def fill_template(
    template_path: Path,
    result: FAResult,
    output_path: Path,
    folder: Path | None = None,
    image_matches: dict[str, list[ImageMatch]] | None = None,
    *,
    update_documents_only: bool = False,
) -> Path:
    template_path = template_path.resolve()
    output_path = output_path.resolve()
    if template_path != output_path:
        shutil.copy2(template_path, output_path)

    support_docs = result.audit_meta.get("support_docs", {}) if result.audit_meta else {}
    support_items: list[tuple[str, Path]] = []
    if isinstance(support_docs, dict):
        for key in ("iso9001", "iso14001", "iso45001", "gmp"):
            path_value = support_docs.get(f"{key}_path")
            if path_value:
                support_items.append((key, Path(str(path_value))))

    with zipfile.ZipFile(output_path, "r") as zin:
        xml_bytes = zin.read("word/document.xml")
        rels_bytes = zin.read("word/_rels/document.xml.rels")
        other_files = {
            name: zin.read(name)
            for name in zin.namelist()
            if name not in ("word/document.xml", "word/_rels/document.xml.rels")
        }

    root = ET.fromstring(xml_bytes)

    if result.factory_info:
        fill_factory_info(root, result.factory_info)
    if result.audit_meta:
        fill_audit_meta(root, result.audit_meta)

    _fill_supporting_doc_numbers(root, folder)

    # Pre-compute RHS text per section in order so document numbers aren't repeated.
    used_doc_nos: set[str] = set()
    rhs_by_sid: dict[str, str] = {}
    for sid in SECTION_IDS:
        if sid in result.sections:
            rhs_by_sid[sid] = _format_section_rhs(
                sid, result.sections[sid], folder=folder, used_doc_nos=used_doc_nos)

    for tbl in root.findall(f".//{W}tbl"):
        for tr in tbl.findall(f"{W}tr"):
            tcs = tr.findall(f"{W}tc")
            if len(tcs) < 3:
                continue
            first = _cell_text(tcs[0]).strip()
            m = re.match(r"^([2-6]\.\d+)", first)
            if not m:
                continue
            sid = m.group(1)
            if sid not in rhs_by_sid:
                continue
            _update_rhs_cell(tcs[2], rhs_by_sid[sid])

    # --- images/signatures/certs (single ET pipeline; always refresh factory photos) ---
    image_matches = image_matches or {}
    media_files: dict[str, bytes] = {}
    signature_payloads: list[tuple[str, int, int, str]] = []
    # section_content maps a FACTORY_PHOTO_LAYOUT source key -> [(rid, cx_emu, cy_emu)]
    section_content: dict[str, list[tuple[int, int, int]]] = {}

    def _uniform_extent(w: int, h: int) -> tuple[int, int]:
        # Force every factory photo to the exact same box so the grid is neat.
        return PHOTO_FIXED_WIDTH_EMU, PHOTO_FIXED_HEIGHT_EMU

    def _fit_extent(w: int, h: int) -> tuple[int, int]:
        # Certificate/master-list pages: keep aspect ratio, cap the height.
        ratio = h / max(w, 1)
        cx = PHOTO_FIXED_WIDTH_EMU
        cy = int(cx * ratio)
        if cy > PHOTO_FIXED_HEIGHT_EMU:
            cy = PHOTO_FIXED_HEIGHT_EMU
            cx = int(cy / max(ratio, 0.01))
        return cx, cy

    flat_images: list[tuple[str, ImageMatch]] = []
    slot_by_id = {s.slot_id: s for s in PHOTO_SLOTS}
    for slot_id, matches in image_matches.items():
        slot = slot_by_id.get(slot_id)
        if not slot or not matches or slot.anchor_type != "paragraph":
            continue
        for im in matches:
            flat_images.append((slot_id, im))

    # General-document images that belong in the Factory Photos section (own doc only)
    doc_image_items: list[tuple[str, Path]] = []
    if folder:
        for key, path in detect_general_doc_images(folder).items():
            doc_image_items.append((key, path))

    signature_items: list[tuple[str, Path]] = []
    if result.audit_meta and result.audit_meta.get("cover_auditor_signature"):
        signature_items.append(("auditor_cover", Path(str(result.audit_meta["cover_auditor_signature"]))))
    for sig_path in result.audit_meta.get("auditor_signatures", []) if result.audit_meta else []:
        signature_items.append(("auditor", Path(sig_path)))
    if result.audit_meta and result.audit_meta.get("factory_signature"):
        signature_items.append(("factory", Path(str(result.audit_meta["factory_signature"]))))

    def _resolve(doc_path: Path) -> Path | None:
        if doc_path.is_file():
            return doc_path
        if folder:
            candidate = folder / doc_path.name
            if candidate.is_file():
                return candidate
            matches = list(folder.rglob(doc_path.name))
            return matches[0] if matches else None
        return None

    doc_page_count = 0
    for key, doc_path in support_items + doc_image_items:
        resolved = _resolve(doc_path)
        if resolved is not None:
            pages = list(_load_document_all_pages(resolved))
            if key == "master_list":
                pages = pages[:1]
            doc_page_count += len(pages)

    total_media = len(flat_images) + doc_page_count + len(signature_items)
    body_el = root.find(f"{W}body")
    if body_el is not None:
        _clean_images_in_section_tables(root)
        _clean_factory_photos_images(root)

    if total_media:
        rel_ids, _ = _next_rel_ids(rels_bytes, total_media)
        rel_items: list[tuple[str, str]] = []
        rel_idx = 0
        photo_idx = _next_fa_media_index(other_files, "fa_")
        cert_idx = _next_fa_media_index(other_files, "fa_cert_")
        sig_idx = _next_fa_media_index(other_files, "fa_signature_")
        for slot_id, im in flat_images:
            raw = _load_image_bytes(Path(im.source_path))
            if not raw:
                continue
            data, ext, w, h = _resize_image(raw[0], raw[1])
            media_name = f"word/media/fa_{photo_idx:03d}.{ext}"
            photo_idx += 1
            media_files[media_name] = data
            rid = rel_ids[rel_idx]
            rel_idx += 1
            rel_items.append((rid, media_name.replace("word/", "")))
            cx, cy = _uniform_extent(w, h)
            section_content.setdefault(slot_id, []).append((rid, cx, cy))

        for key, doc_path in support_items + doc_image_items:
            resolved = _resolve(doc_path)
            if resolved is None:
                continue
            pages = _load_document_all_pages(resolved)
            if key == "master_list":
                pages = pages[:1]
            for page_data, page_ext in pages:
                data, ext, w, h = _resize_image(page_data, page_ext, force_jpeg=True)
                media_name = f"word/media/fa_cert_{cert_idx:03d}.{ext}"
                cert_idx += 1
                media_files[media_name] = data
                rid = rel_ids[rel_idx]
                rel_idx += 1
                rel_items.append((rid, media_name.replace("word/", "")))
                cx, cy = _fit_extent(w, h)
                section_content.setdefault(key, []).append((rid, cx, cy))

        for kind, sig_path in signature_items:
            raw = _load_image_bytes(sig_path)
            if not raw:
                continue
            data, ext, w, h = _resize_image(raw[0], raw[1])
            media_name = f"word/media/fa_signature_{sig_idx:03d}.{ext}"
            sig_idx += 1
            media_files[media_name] = data
            rid = rel_ids[rel_idx]
            rel_idx += 1
            rel_items.append((rid, media_name.replace("word/", "")))
            signature_payloads.append((kind, w, h, rid))

        if rel_items:
            rels_bytes = _add_image_relationships(rels_bytes, rel_items)

        # Deterministic slot-driven rebuild of the entire Factory Photos section
        _build_factory_photos_section(root, section_content)

        if signature_payloads:
            _insert_signatures_in_document(
                root, signature_payloads, skip_if_present=update_documents_only,
            )

    new_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)

    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in other_files.items():
            zout.writestr(name, data)
        zout.writestr("word/document.xml", new_xml)
        zout.writestr("word/_rels/document.xml.rels", rels_bytes)
        for name, data in media_files.items():
            zout.writestr(name, data)

    return output_path


def write_summary(result: FAResult, path: Path) -> None:
    lines = [
        "FA DOCUMENT MATCH SUMMARY",
        "=" * 40,
        f"Folder: {result.folder}",
        f"Master list: {result.master_list or 'NOT FOUND'}",
        "",
    ]
    for sid in SECTION_IDS:
        matches = result.sections.get(sid, [])
        lines.append(f"{sid}")
        if not matches:
            lines.append("  NOT FOUND IN UPLOADED FILES OR MASTER LIST")
        else:
            for m in matches:
                doc = f"{m.doc_no} — " if m.doc_no else ""
                if m.status == "master_list_only":
                    status = "listed in master list only"
                else:
                    status = "provided file matched"
                lines.append(f"  • {doc}{m.heading} [{m.filename}] ({status})")
        lines.append("")
    if result.images:
        lines.append("IMAGES")
        lines.append("-" * 20)
        for slot_id, matches in result.images.items():
            for m in matches:
                lines.append(f"  {m.slot_label}: {m.filename}")
        lines.append("")
    if result.checklist:
        provided = sum(1 for c in result.checklist if c.status == "provided_file")
        master_only = sum(1 for c in result.checklist if c.status == "master_list_only")
        missing = sum(1 for c in result.checklist if c.status == "not_found")
        lines.append(
            "CHECKLIST: "
            f"{provided} provided file matched, "
            f"{master_only} listed in master list only, "
            f"{missing} not found"
        )
        lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")


def merge_supplement_folder(primary_folder: Path, supplement_folder: Path) -> list[str]:
    """Copy supplement files into primary/_supplement/ so matching picks them up."""
    primary_folder = primary_folder.resolve()
    supplement_folder = supplement_folder.resolve()
    dest_root = primary_folder / "_supplement"
    dest_root.mkdir(parents=True, exist_ok=True)
    copied: list[str] = []

    for path in supplement_folder.rglob("*"):
        if not path.is_file():
            continue
        if path.name.startswith(".") or path.name.lower() in {".ds_store", "thumbs.db"}:
            continue
        if path.suffix.lower() not in {
            ".pdf", ".xlsx", ".xls", ".docx", ".doc", ".jpeg", ".jpg", ".png", ".webp"
        }:
            continue
        dest = dest_root / path.name
        if dest.exists():
            stem, suffix = dest.stem, dest.suffix
            n = 2
            while dest.exists():
                dest = dest_root / f"{stem}_{n}{suffix}"
                n += 1
        shutil.copy2(path, dest)
        copied.append(dest.name)

    return sorted(copied)


def diff_fa_results(before: FAResult | dict[str, Any], after: FAResult | dict[str, Any]) -> dict[str, Any]:
    """Summarize what changed after merging a supplement folder."""

    def _sections(obj: FAResult | dict[str, Any]) -> dict[str, list[dict[str, Any]]]:
        if isinstance(obj, FAResult):
            return obj.to_dict()["sections"]
        return obj.get("sections", {})

    def _checklist(obj: FAResult | dict[str, Any]) -> list[dict[str, Any]]:
        if isinstance(obj, FAResult):
            return obj.to_dict()["checklist"]
        return obj.get("checklist", [])

    before_sections = _sections(before)
    after_sections = _sections(after)
    before_checklist = _checklist(before)
    after_checklist = _checklist(after)

    def _doc_key(m: dict[str, Any]) -> str:
        return f"{m.get('doc_no', '')}|{m.get('filename', '')}|{m.get('heading', '')}".lower()

    new_docs: dict[str, list[dict[str, Any]]] = {}
    for sid in SECTION_IDS:
        seen = {_doc_key(m) for m in before_sections.get(sid, [])}
        added = [m for m in after_sections.get(sid, []) if _doc_key(m) not in seen]
        if added:
            new_docs[sid] = added

    before_missing = {c["label"] for c in before_checklist if not c.get("present")}
    after_missing = {c["label"] for c in after_checklist if not c.get("present")}
    resolved = sorted(before_missing - after_missing)

    return {
        "new_docs": new_docs,
        "resolved_checklist": resolved,
        "missing_before": len(before_missing),
        "missing_after": len(after_missing),
    }


def run_fa_supplement(
    primary_folder: Path,
    supplement_folder: Path,
    output_dir: Path | None = None,
    template_path: Path | None = None,
    filled_docx_path: Path | None = None,
    auditor_names: list[str] | None = None,
) -> tuple[FAResult, list[str]]:
    """Merge supplement files, rebuild checklist, update the already-filled FA report."""
    primary_folder = primary_folder.resolve()
    output_dir = (output_dir or primary_folder).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    copied = merge_supplement_folder(primary_folder, supplement_folder)

    result = match_documents(primary_folder)
    image_matches = match_images(primary_folder)
    result.images = image_matches
    result.factory_info = parse_factory_info(primary_folder)
    result.audit_meta = extract_audit_meta(primary_folder, result.factory_info, auditor_names)

    master_path = Path(result.master_list) if result.master_list else None
    checklist_items = build_checklist(primary_folder, image_matches, master_path)
    result.checklist = checklist_items

    factory_name = primary_folder.name[:40].replace("/", "-")
    checklist_path = output_dir / f"FA_Checklist_{factory_name}.txt"
    write_checklist(checklist_items, checklist_path)
    result.checklist_path = str(checklist_path)

    filled_docx = Path(filled_docx_path).resolve() if filled_docx_path else None
    if filled_docx and filled_docx.exists():
        out_docx = filled_docx
        fill_template(
            filled_docx,
            result,
            out_docx,
            folder=primary_folder,
            image_matches=image_matches,
            update_documents_only=True,
        )
    else:
        template_path = (template_path or DEFAULT_TEMPLATE).resolve()
        if not template_path.exists():
            raise FileNotFoundError(f"FA template not found: {template_path}")
        out_docx = output_dir / f"FA_filled_{factory_name}.docx"
        fill_template(
            template_path,
            result,
            out_docx,
            folder=primary_folder,
            image_matches=image_matches,
        )

    summary_path = output_dir / f"FA_summary_{factory_name}.txt"
    write_summary(result, summary_path)

    result.output_docx = str(out_docx)
    result.summary_path = str(summary_path)

    json_path = output_dir / f"FA_summary_{factory_name}.json"
    json_path.write_text(json.dumps(result.to_dict(), indent=2), encoding="utf-8")

    return result, copied


def run_fa(
    folder: Path,
    output_dir: Path | None = None,
    template_path: Path | None = None,
    auditor_names: list[str] | None = None,
) -> FAResult:
    folder = folder.resolve()
    template_path = (template_path or DEFAULT_TEMPLATE).resolve()
    if not template_path.exists():
        raise FileNotFoundError(f"FA template not found: {template_path}")

    output_dir = (output_dir or folder).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    result = match_documents(folder)
    image_matches = match_images(folder)
    result.images = image_matches
    result.factory_info = parse_factory_info(folder)
    result.audit_meta = extract_audit_meta(folder, result.factory_info, auditor_names)

    master_path = Path(result.master_list) if result.master_list else None
    checklist_items = build_checklist(folder, image_matches, master_path)
    result.checklist = checklist_items

    factory_name = folder.name[:40].replace("/", "-")
    checklist_path = output_dir / f"FA_Checklist_{factory_name}.txt"
    write_checklist(checklist_items, checklist_path)
    result.checklist_path = str(checklist_path)

    out_docx = output_dir / f"FA_filled_{factory_name}.docx"
    fill_template(template_path, result, out_docx, folder=folder, image_matches=image_matches)

    summary_path = output_dir / f"FA_summary_{factory_name}.txt"
    write_summary(result, summary_path)

    result.output_docx = str(out_docx)
    result.summary_path = str(summary_path)

    json_path = output_dir / f"FA_summary_{factory_name}.json"
    json_path.write_text(json.dumps(result.to_dict(), indent=2), encoding="utf-8")

    return result


def main():
    import argparse

    ap = argparse.ArgumentParser(description="FA document matcher + template filler")
    ap.add_argument("folder", type=Path, help="Folder with client FA documents")
    ap.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    ap.add_argument("--output", type=Path, default=None)
    args = ap.parse_args()

    result = run_fa(args.folder, args.output, args.template)
    print(f"Filled template: {result.output_docx}")
    print(f"Summary: {result.summary_path}")
    for sid in SECTION_IDS:
        n = len(result.sections.get(sid, []))
        print(f"  {sid}: {n} document(s)")


if __name__ == "__main__":
    main()
